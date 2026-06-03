from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path("deliverables/Bo_cau_hoi_phan_bien_DA_CHINH_SUA.docx")


QA_ITEMS = [
    (
        "1. Rò rỉ dữ liệu ở cấp độ bệnh nhân",
        "\"Trong slide 3 và slide 6, nhóm tự tin khẳng định dùng giao thức chia dữ liệu theo lesion_id để chống rò rỉ dữ liệu (Data Leakage). Tuy nhiên, về mặt y khoa, một bệnh nhân có thể có nhiều tổn thương với các mã lesion_id khác nhau trên cùng một cơ thể. Chia theo lesion_id hoàn toàn không ngăn chặn được việc ảnh của cùng một bệnh nhân xuất hiện ở cả tập Train và Test. Tại sao nhóm lại bỏ qua rủi ro Patient-level Leakage này?\"",
        [
            "Nhóm không phủ nhận rủi ro này. Giao thức hiện tại kiểm soát rò rỉ ở cấp độ tổn thương: mọi ảnh có cùng lesion_id được giữ trong cùng một phần dữ liệu. Điều đó ngăn một tổn thương xuất hiện đồng thời ở tập học và tập đánh giá, nhưng không tương đương với patient-level split.",
            "Trong các tệp metadata đang sử dụng cho HAM10000 và BCN20000 của dự án, nhóm có trường lesion_id nhưng không có patient_id. Vì vậy, với dữ liệu hiện có, nhóm chưa thể chứng minh rằng hai tổn thương khác nhau của cùng bệnh nhân không bị phân bố sang hai tập khác nhau.",
            "Kết luận đúng là: nhóm đã giảm một nguồn leakage quan trọng ở cấp lesion, nhưng patient-level leakage vẫn là giới hạn dữ liệu cần công bố minh bạch. Nếu nhận được patient_id hoặc bộ dữ liệu có định danh bệnh nhân, giao thức ưu tiên phải là patient-level group split. Kết quả cross-dataset chỉ được dùng như đánh giá khả năng tổng quát hóa khi đã chạy và báo cáo riêng, không được dùng để tuyên bố loại bỏ patient-level leakage.",
        ],
    ),
    (
        "2. FiLM phụ thuộc vào backbone",
        "\"Hãy nhìn vào bảng kết quả trên tập HAM10000: Khi cùng tích hợp cơ chế FiLM (Strategy 3), ConvNeXt đạt 94.02% AUC, nhưng EfficientNet-B4 chỉ đạt 85.44% - sụt giảm tới gần 9%! Tại sao cơ chế điều biến của nhóm lại mất ổn định và nhạy cảm cực đoan theo kiến trúc mạng như vậy? Làm sao đảm bảo tính an toàn khi triển khai thực tế?\"",
        [
            "Hai con số trên cho thấy hiệu quả của FiLM phụ thuộc vào backbone; chúng không đủ để kết luận nguyên nhân là một cơ chế kiến trúc cụ thể như depthwise convolution hay compound scaling. FiLM trong dự án sinh gamma và beta từ metadata rồi điều biến feature ảnh, nên mức hữu ích của metadata phụ thuộc vào biểu diễn mà từng backbone tạo ra và cách tối ưu hội tụ.",
            "Vì vậy, nhóm không khẳng định FiLM luôn tốt hơn hoặc ổn định trên mọi kiến trúc. Kết luận có thể bảo vệ là: FiLM là một cơ chế fusion cần được thẩm định theo từng backbone và từng tập dữ liệu; cấu hình triển khai chỉ được chọn sau khi đánh giá trên protocol cố định và báo cáo đầy đủ AUC, sensitivity, specificity và độ biến thiên.",
            "Trong artefact EfficientNet-B4 hiện có của dự án, AUC HAM của image-only là 85.37%, FiLM là 85.43%, FiLM weighted là 85.44%, còn gating là 82.54%. Kiểm định AUC đang lưu không cho thấy FiLM khác image-only có ý nghĩa thống kê. Do đó, nhóm không sử dụng EfficientNet-B4 để khẳng định ưu thế phổ quát của FiLM.",
        ],
    ),
    (
        "3. Báo cáo đầy đủ bốn chiến lược fusion",
        "\"Trong báo cáo và slide 4, nhóm dành rất nhiều thời gian phân tích và đề xuất đến 4 chiến lược dung hợp (Baseline, Concatenation, FiLM, Gating). Thế nhưng tại sao trong bảng kết quả thực nghiệm, nhóm chỉ đưa ra duy nhất kết quả của Strategy 3 (FiLM)? Kết quả của hai chiến lược Concatenation và Gating đâu? Phải chăng chúng quá tệ nên nhóm đã giấu đi?\"",
        [
            "Nếu một slide chỉ trình bày FiLM mà không kèm bảng đối chứng, đó là thiếu sót trình bày và cần sửa. Trong code, nhóm có triển khai image-only, concatenation, FiLM và gating; vì vậy báo cáo khoa học phải thể hiện các kết quả đối chứng thay vì chỉ nêu cấu hình tốt nhất.",
            "Với các checkpoint EfficientNet-B4 hiện đang truy vết được, HAM10000 cho AUC lần lượt: image-only 85.37%, FiLM 85.43%, FiLM weighted 85.44% và gating 82.54%. Trên BCN20000, các giá trị tương ứng là 79.31%, 79.31%, 79.25% và 79.22%. Các số liệu này không hỗ trợ tuyên bố rằng gating tương đương FiLM hoặc FiLM luôn vượt trội.",
            "Đối với Concatenation hoặc các backbone được nêu trong bảng của bài báo nhưng không có đầy đủ artefact tái lập trong workspace hiện tại, nhóm sẽ đính kèm bảng gốc có nguồn rõ ràng hoặc chạy lại cùng protocol trước khi đưa vào kết luận. Cách trả lời đúng là công khai toàn bộ so sánh, không giải thích việc thiếu bảng bằng suy đoán về chi phí hay thời gian hội tụ.",
        ],
    ),
    (
        "4. Sensitivity trong bài toán sàng lọc",
        "\"Tại sao trong bảng kết quả thực nghiệm, nhóm chỉ báo cáo hai chỉ số là AUC và Độ đặc hiệu (Specificity)? Trong ung thư da, bỏ sót một ca ác tính, tức Âm tính giả, mới là nguy cơ nghiêm trọng. Việc nhóm không báo cáo Độ nhạy có phải là để che giấu việc mô hình đang bỏ sót hàng loạt ca ung thư?\"",
        [
            "Hội đồng đặt đúng trọng tâm: trong ứng dụng sàng lọc, sensitivity/recall phải được báo cáo cùng AUC và specificity. Code đánh giá của nhóm đã tính recall, precision và F1, do đó việc không đưa sensitivity lên slide là thiếu sót báo cáo chứ không phải một chỉ số không có sẵn.",
            "Quan trọng hơn, kết quả EfficientNet-B4 hiện có cho HAM10000 tại threshold 0.5 cho recall khá thấp: image-only 36.83%, FiLM 40.40%, FiLM weighted 37.52% và gating 23.66%. Trên BCN20000, recall lần lượt là 68.75%, 70.35%, 70.18% và 68.41%. Nhóm không thể dùng AUC hoặc specificity để khẳng định hệ thống đã an toàn cho sàng lọc lâm sàng.",
            "Kết luận phù hợp là mô hình hiện mới ở mức nghiên cứu phân loại nguy cơ. Trước khi hướng tới sử dụng lâm sàng, nhóm phải lựa chọn threshold theo mục tiêu sensitivity, báo cáo confusion matrix và khoảng tin cậy, đồng thời đánh giá đánh đổi giữa giảm false negative và tăng false positive trên tập kiểm định độc lập.",
        ],
    ),
    (
        "5. DullRazor, CLAHE và thông tin màu sắc",
        "\"Nhóm sử dụng thuật toán DullRazor để xóa lông và thuật toán CLAHE để tăng cường tương phản. Tuy nhiên, trong y khoa, quy trình ABCD để chẩn đoán u hắc tố dựa rất mạnh vào yếu tố C (Color). Thuật toán CLAHE làm biến đổi cục bộ biểu đồ màu sắc và độ sáng của bức ảnh. Chẳng phải nhóm đang tự tay phá hủy các đặc trưng màu sắc sinh học cốt lõi, ép AI học trên một bức ảnh bị bóp méo hay sao?\"",
        [
            "Đây là rủi ro hợp lệ và nhóm không khẳng định ảnh sau tiền xử lý giữ nguyên màu theo nghĩa pixel-by-pixel. Trong code, bước CLAHE được thực hiện sau khi chuyển ảnh sang LAB và chỉ áp dụng trên kênh L với clipLimit bằng 1.5; hai kênh sắc độ a và b không bị CLAHE biến đổi trực tiếp. Mục tiêu của bước này là giảm biến thiên độ sáng và làm rõ cấu trúc quan sát được.",
            "Tuy nhiên, pipeline đầy đủ còn có Gray-World trước CLAHE để hiệu chỉnh cân bằng màu, inpainting tại vùng được nhận diện là lông, và bilateral filtering sau khi tăng tương phản. Ngoài ra, tập train có ColorJitter. Vì vậy, câu nói rằng các đặc trưng màu sinh học được bảo toàn nguyên vẹn là không chính xác.",
            "Lập luận khoa học đúng là pipeline cố gắng giảm nhiễu do lông và chiếu sáng trong khi hạn chế can thiệp trực tiếp vào sắc độ ở bước CLAHE. Để xác nhận lựa chọn này không làm mất tín hiệu chẩn đoán, nhóm cần ablation ảnh gốc, xóa lông, Gray-World và CLAHE; đồng thời so sánh sensitivity/AUC, đặc biệt với melanoma, trước khi kết luận lợi ích của tiền xử lý.",
        ],
    ),
    (
        "6. Giới hạn của Grad-CAM và SHAP",
        "\"Nhóm sử dụng Grad-CAM để chứng minh AI đã 'nhìn' vào đúng chỗ tổn thương nhằm xây dựng niềm tin lâm sàng. Tuy nhiên, Grad-CAM có thể bị ảnh hưởng bởi vùng gradient mạnh hoặc tương phản do CLAHE tạo ra. Làm sao nhóm khẳng định vùng màu đỏ thực sự là tri thức y khoa?\"",
        [
            "Nhóm không nên dùng từ 'chứng minh' cho Grad-CAM. Grad-CAM là công cụ trực quan hóa vùng ảnh ảnh hưởng tới dự đoán, không phải xác nhận rằng mô hình đã học đúng đặc trưng y khoa hoặc rằng vùng chú ý không chịu ảnh hưởng của preprocessing.",
            "Trong triển khai hiện tại, hàm sinh Grad-CAM đưa metadata cố định bằng tensor zero khi gọi mô hình. Vì vậy, các heatmap hiện tại chủ yếu khảo sát nhánh ảnh dưới một điều kiện metadata cố định; chúng chưa chứng minh sự tương tác động giữa ảnh và metadata. Các script SHAP cũng là phân tích đóng góp metadata, không phải bằng chứng nhân quả về vị trí tổn thương.",
            "Cách trả lời phù hợp là: Grad-CAM và SHAP được sử dụng như phân tích hỗ trợ. Để kiểm tra nghiêm ngặt lo ngại của hội đồng, nhóm cần bổ sung so sánh heatmap trên ảnh gốc và ảnh tiền xử lý, kiểm thử occlusion/perturbation định lượng, và nếu đánh giá multimodal saliency thì phải giữ metadata thực hoặc thay đổi metadata có kiểm soát.",
        ],
    ),
    (
        "7. Khả năng triển khai trên thiết bị di động",
        "\"Nhóm định hướng nén mô hình đưa lên ứng dụng di động, đồng thời tích hợp LLM để đọc EHR. Một Vision Transformer kết hợp LLM đòi hỏi tài nguyên lớn. Làm sao chạy trên điện thoại cấu hình thấp ở trạm y tế xã?\"",
        [
            "Đây là định hướng tương lai, chưa phải năng lực đã được triển khai hoặc benchmark trong dự án hiện tại. Nhóm không chủ trương chạy đồng thời một ViT lớn và một LLM y tế hoàn chỉnh offline trên điện thoại cấu hình thấp.",
            "Kiến trúc khả thi về mặt thiết kế là phân tách nhiệm vụ: thiết bị đầu cuối có thể chạy một bộ mã hóa ảnh nhỏ sau khi đã được lựa chọn, lượng tử hóa hoặc distillation; tác vụ xử lý văn bản EHR phức tạp chỉ thực hiện trên máy chủ khi có hạ tầng mạng, kiểm soát truy cập và bảo mật dữ liệu phù hợp.",
            "Tuy nhiên, để biến định hướng thành kết quả nghiên cứu, nhóm phải báo cáo kích thước model, latency, bộ nhớ, mức giảm hiệu năng sau nén, khả năng hoạt động offline và quy trình bảo mật. Hiện tại nhóm chỉ trình bày đây là lộ trình triển khai, không phải kết quả thực nghiệm đã chứng minh.",
        ],
    ),
    (
        "8. Ý nghĩa khoa học của mức tăng AUC",
        "\"Strategy 3 (FiLM) đạt AUC 94.02%, cao hơn Baseline khoảng 1.5 - 2%. Làm sao nhóm chứng minh đây là sự vượt trội có ý nghĩa khoa học chứ không phải do nhiễu ngẫu nhiên hoặc khởi tạo trọng số?\"",
        [
            "Nếu con số 94.02% được lấy từ bảng ConvNeXt của bài báo, nhóm chỉ có thể nói đó là kết quả trung bình theo các fold được báo cáo trong bảng tương ứng. Không được tự động diễn giải thành kết quả của nhiều random seed nếu chưa có thí nghiệm seed robustness riêng.",
            "Trong workspace hiện tại, kiểm định paired theo 5 fold đã được lưu cho EfficientNet-B4. Trên HAM10000, FiLM weighted chỉ cao hơn image-only khoảng 0.075 điểm phần trăm về AUC với p = 0.8531; trên BCN20000, FiLM weighted thấp hơn image-only khoảng 0.062 điểm phần trăm với p = 0.8920. Những kết quả này không cho phép tuyên bố FiLM vượt trội có ý nghĩa thống kê trên EfficientNet-B4.",
            "Để bảo vệ tuyên bố cho ConvNeXt 94.02%, nhóm cần cung cấp dự đoán theo fold hoặc theo mẫu của chính cấu hình ConvNeXt, thực hiện kiểm định phù hợp và báo cáo khoảng tin cậy. Cho đến khi có kiểm định tương ứng, kết luận nên giới hạn ở mức: FiLM cho kết quả triển vọng ở một số backbone trong bảng báo cáo, nhưng mức cải thiện chưa được chứng minh là phổ quát.",
        ],
    ),
    (
        "9. Focal Loss và hiệu chuẩn xác suất",
        "\"Nhóm dùng Focal Loss để xử lý mất cân bằng. Hàm này có thể làm sai lệch confidence score. Dù Accuracy cao, độ tin cậy của xác suất đưa cho bác sĩ có thể bị sai. Nhóm đã đo lường ECE chưa?\"",
        [
            "Nhóm thừa nhận điểm này. Dự án hiện sử dụng Focal BCE Loss với alpha bằng 0.75 và gamma bằng 2.0 để tập trung học các mẫu khó, nhưng trong code và artefact hiện có chưa có đánh giá ECE, Brier score hoặc calibration curve.",
            "Do đó, đầu ra sigmoid hiện tại không nên được mô tả như một xác suất nguy cơ đã được hiệu chuẩn để bác sĩ sử dụng trực tiếp. AUC đánh giá khả năng xếp hạng và các chỉ số tại threshold đánh giá quyết định phân lớp; chúng không chứng minh confidence score đáng tin cậy.",
            "Bước bổ sung bắt buộc trước khi trình bày xác suất lâm sàng là hiệu chuẩn trên tập validation tách biệt, chẳng hạn temperature scaling, sau đó báo cáo ECE, Brier score, reliability diagram và đánh giá lại sensitivity/specificity tại threshold được chọn.",
        ],
    ),
    (
        "10. Lợi ích metadata hay chỉ do tăng số tham số",
        "\"FiLM cho kết quả cao hơn Baseline, nhưng FiLM cũng thêm tham số. Làm sao nhóm biết mức tăng đến từ thông tin metadata thay vì capacity lớn hơn giúp mô hình overfit?\"",
        [
            "So sánh image-only với FiLM có metadata thật chưa đủ để tách riêng hai giả thuyết: lợi ích của thông tin lâm sàng và lợi ích của việc tăng số tham số. Trong workspace hiện tại, nhóm chưa có artefact chứng minh đã chạy thí nghiệm shuffled metadata hoặc dummy metadata.",
            "Thí nghiệm đúng để trả lời câu hỏi là giữ nguyên kiến trúc FiLM và số tham số, nhưng huấn luyện/đánh giá với metadata bị hoán vị giữa các mẫu hoặc metadata hằng số. Nếu FiLM với metadata thật vượt rõ rệt các đối chứng cùng capacity, khi đó mới có bằng chứng rằng thông tin lâm sàng đóng góp thực sự.",
            "Vì chưa có ablation đó trong kết quả hiện tại, nhóm chỉ kết luận rằng mô hình đã triển khai cơ chế điều biến bằng metadata và có thể khảo sát đóng góp feature; nhóm chưa tuyên bố mức tăng AUC, nếu có, hoàn toàn do ý nghĩa y khoa của metadata.",
        ],
    ),
    (
        "11. Domain shift và phép so sánh giữa HAM10000/BCN20000",
        "\"Mô hình giảm từ khoảng 94% trên HAM10000 xuống 84.88% trên BCN20000. Việc giảm đến 10% chứng tỏ mô hình bị Domain Shift rất nặng. Mô hình có thực sự học được khái quát sinh học không?\"",
        [
            "Cần chỉnh lại tiền đề của phép so sánh. Hai số 94.02% trên HAM10000 và 84.88% trên BCN20000 trong tài liệu trình bày gắn với các cấu hình/backbone báo cáo khác nhau; so sánh trực tiếp chúng không phải là một phép đo domain shift hợp lệ.",
            "Đánh giá domain shift đúng phải cố định mô hình và hướng chuyển miền: ví dụ train trên HAM rồi test trực tiếp trên BCN, hoặc ngược lại, với cùng checkpoint, cùng ánh xạ metadata và cùng chỉ số. Dự án có script cho đánh giá chéo, nhưng trong artefact hiện đang đối chiếu chưa có bảng kết quả cross-dataset được lưu để dùng làm bằng chứng định lượng.",
            "Vì vậy, nhóm chỉ có thể nói hai bộ dữ liệu có khác biệt phân phối và việc đánh giá chéo là cần thiết. Nhóm chưa dùng chênh lệch giữa hai bảng nội miền để tuyên bố mô hình đã kháng domain shift hoặc đã học được khái quát sinh học.",
        ],
    ),
    (
        "12. Biến metadata nào đóng góp chính",
        "\"Nhóm nạp Tuổi, Giới tính, Vị trí vào mô hình. Trong mức tăng hiệu suất đó, biến nào đóng vai trò chính? Nếu loại bỏ Giới tính, mô hình có bị ảnh hưởng không?\"",
        [
            "Trong dữ liệu HAM, nhánh metadata sử dụng age, sex và localization; trong BCN sử dụng age_approx, sex và anatom_site_general. Các bảng importance hiện có cho thấy tuổi là feature có trọng số giải thích lớn nhất: khoảng 0.630 trong artefact HAM và khoảng 0.807 trong artefact BCN; sau đó là các nhóm vị trí giải phẫu, còn các biến giới tính có importance nhỏ hơn.",
            "Tuy nhiên, feature importance hoặc SHAP chỉ phản ánh mức đóng góp trong mô hình/phân tích hiện tại; chúng không tương đương với bằng chứng rằng bỏ giới tính chỉ làm AUC thay đổi một con số xác định. Muốn trả lời định lượng, nhóm phải chạy ablation loại riêng age, sex và location trên cùng protocol.",
            "Do đó, kết luận được phép nêu là: phân tích hiện có gợi ý tuổi là tín hiệu metadata nổi bật nhất và giới tính có đóng góp nhỏ hơn trong các artefact đang lưu. Nhóm chưa khẳng định tác động nhân quả của từng biến lên AUC khi chưa có thí nghiệm loại biến.",
        ],
    ),
    (
        "13. Thiếu hoặc nhập sai metadata khi triển khai",
        "\"FiLM giúp mô hình chú ý metadata, nhưng nếu bác sĩ nhập sai tuổi hoặc quên nhập vị trí thì sao? Mô hình có phụ thuộc quá mức vào metadata và sụp đổ so với image-only không?\"",
        [
            "Đây là rủi ro triển khai cần đánh giá riêng. Code hiện tại hỗ trợ giá trị phân loại unknown trong xử lý dữ liệu và FiLM được khởi tạo ban đầu gần identity do lớp sinh gamma/beta khởi tạo zero. Tuy nhiên, sau huấn luyện, điều đó không bảo đảm mô hình tự động trở về image-only khi metadata thiếu hoặc sai.",
            "Trong artefact hiện có, nhóm chưa có robustness test che metadata, nhiễu tuổi, hoán đổi vị trí hoặc đo mức suy giảm hiệu năng khi nhập sai. Vì vậy, nhóm không tuyên bố mô hình có cơ chế safe-fail đã được chứng minh.",
            "Trước triển khai, nhóm cần bổ sung missing-modality/mis-entry testing, huấn luyện với metadata dropout hoặc missingness mask, và cân nhắc cung cấp một nhánh image-only fallback được đánh giá độc lập. Khi đó mới có thể quy định rõ hệ thống xử lý trường hợp metadata không đáng tin cậy như thế nào.",
        ],
    ),
]


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(document, text, *, size=11, bold=False, color=None, space_after=4):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BỘ CÂU HỎI PHẢN BIỆN VÀ TRẢ LỜI ĐÃ HIỆU CHỈNH")
    set_font(run, size=15, bold=True, color=(31, 78, 121))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Đối chiếu với code và artefact hiện có của dự án chẩn đoán ung thư da đa phương thức")
    set_font(run, size=11, color=(89, 89, 89))

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(12)
    note.paragraph_format.line_spacing = 1.15
    run = note.add_run(
        "Nguyên tắc hiệu chỉnh: chỉ phát biểu các kết luận có thể kiểm chứng từ code hoặc kết quả đang lưu; "
        "các thí nghiệm chưa có artefact được nêu rõ là giới hạn hoặc công việc cần bổ sung."
    )
    set_font(run, size=10, bold=True, color=(192, 0, 0))

    for title_text, question, answers in QA_ITEMS:
        add_paragraph(document, title_text, size=12, bold=True, color=(31, 78, 121), space_after=3)
        add_paragraph(document, question, size=11, bold=True, space_after=5)
        add_paragraph(document, "Trả lời đề xuất:", size=11, bold=True, color=(0, 97, 0), space_after=3)
        for answer in answers:
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.first_line_indent = Cm(-0.35)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.15
            bullet = paragraph.add_run("- ")
            set_font(bullet, size=11, bold=True)
            run = paragraph.add_run(answer)
            set_font(run, size=11)
        document.add_paragraph()

    add_paragraph(document, "Căn cứ đối chiếu chính trong workspace", size=12, bold=True, color=(31, 78, 121))
    references = [
        "src/preprocessed/preprocess_pipeline.py: DullRazor-style inpainting, Gray-World, CLAHE trên kênh L của LAB và bilateral filtering.",
        "src/data_logic/common_transforms.py: ColorJitter chỉ áp dụng trong augmentation của tập train.",
        "src/utils/experiment_runner.py và script/create_group_splits.py: protocol nhóm theo lesion_id.",
        "src/models/fusion_head.py và src/models/__init__.py: triển khai Concatenation, FiLM và Gating.",
        "src/utils/trainer.py: đánh giá AUC, accuracy, F1, precision, recall và specificity; Grad-CAM sử dụng metadata cố định bằng zero tensor.",
        "results/significance_tests_auc.csv và các tệp summary trong checkpoint_ham10000/checkpoint_bcn20000: số liệu EfficientNet-B4 và kiểm định AUC hiện có.",
    ]
    for reference in references:
        add_paragraph(document, "- " + reference, size=10, space_after=3)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH.resolve())
