from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parents[2] / "deliverables"
OUT_FILE = OUT_DIR / "Danh_sach_y_trong_tam_va_cau_hoi_hoi_dong_NCKH_ung_thu_da_DA_HIEU_DINH.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(9)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "1F4E79")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_qa(doc, question, answer, caution=None):
    p = doc.add_paragraph()
    q = p.add_run("CÃ¢u há»i: ")
    q.bold = True
    p.add_run(question)
    p = doc.add_paragraph()
    a = p.add_run("Tráº£ lá»i an toÃ n: ")
    a.bold = True
    p.add_run(answer)
    if caution:
        p = doc.add_paragraph()
        c = p.add_run("LÆ°u Ã½: ")
        c.bold = True
        c.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        p.add_run(caution)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        styles[style_name].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DANH SÃCH Ã TRá»ŒNG TÃ‚M VÃ€ CÃ‚U Há»ŽI Há»˜I Äá»’NG\n"
                        "Báº¢N HIá»†U ÄÃNH THEO BÃ€I BÃO, SLIDE, SCRIPT VÃ€ MÃƒ NGUá»’N")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Äá» tÃ i: Multimodal Deep Learning with Clinical Metadata for Enhanced Skin Cancer Diagnosis\n"
        "NgÃ y Ä‘á»‘i chiáº¿u: 25/05/2026 | Tráº¡ng thÃ¡i: Báº£n sá»­a phá»¥c vá»¥ chuáº©n bá»‹ pháº£n biá»‡n"
    ).italic = True

    p = doc.add_paragraph()
    r = p.add_run("Káº¿t luáº­n ngáº¯n: ")
    r.bold = True
    p.add_run(
        "BÃ i bÃ¡o CITA bá»• sung lÃ  nguá»“n cá»§a cÃ¡c báº£ng káº¿t quáº£ bá»‘n backbone vÃ  bá»‘n fusion strategy Ä‘Æ°á»£c dÃ¹ng "
        "trÃªn slide, bao gá»“m ConvNeXt/ViT. Tuy nhiÃªn, khi Ä‘á»‘i chiáº¿u vá»›i mÃ£ nguá»“n/checkpoint trong workspace, "
        "chá»‰ má»™t pháº§n káº¿t quáº£ EfficientNet-B4 cÃ³ thá»ƒ kiá»ƒm tra trá»±c tiáº¿p vÃ  cÃ²n cÃ¡c sai khÃ¡c quan trá»ng. "
        "Äáº·c biá»‡t, code BCN váº«n Ä‘Æ°a `diagnosis_confirm_type` vÃ o mÃ´ hÃ¬nh dÃ¹ paper tuyÃªn bá»‘ chá»‰ dÃ¹ng metadata "
        "trÆ°á»›c cháº©n Ä‘oÃ¡n. VÃ¬ váº­y, khi báº£o vá»‡ cáº§n phÃ¢n biá»‡t rÃµ káº¿t quáº£ cÃ´ng bá»‘ trong paper vá»›i báº±ng chá»©ng "
        "tÃ¡i láº­p hiá»‡n cÃ³ vÃ  khÃ´ng dÃ¹ng cÃ¡c kháº³ng Ä‘á»‹nh tuyá»‡t Ä‘á»‘i vá» leakage hay Ã½ nghÄ©a lÃ¢m sÃ ng."
    )

    heading(doc, "1. Nguá»“n Ä‘Ã£ Ä‘á»‘i chiáº¿u", 1)
    sources = [
        "BÃ i bÃ¡o: CITA_2026_paper_357.pdf, 14 trang; nguá»“n chÃ­nh thá»©c cá»§a Table 1, Table 2 vÃ  mÃ´ táº£ phÆ°Æ¡ng phÃ¡p cÃ´ng bá»‘.",
        "Slide: BÃ¡o cÃ¡o NCKH 2025 - 2026.pdf, 15 trang.",
        "Script thuyáº¿t trÃ¬nh: Script bÃ¡o cÃ¡o.docx.",
        "TÃ i liá»‡u cáº§n hiá»‡u Ä‘Ã­nh: Danh_sach_y_trong_tam_va_cau_hoi_hoi_dong_NCKH_ung_thu_da.docx.",
        "MÃ£ huáº¥n luyá»‡n: script/training/train_ham.py vÃ  script/training/train_bcn.py.",
        "Kiáº¿n trÃºc: src/models/fusion_head.py, src/models/__init__.py vÃ  cÃ¡c backbone.",
        "Dá»¯ liá»‡u/tiá»n xá»­ lÃ½: src/data_logic/*.py vÃ  src/preprocessed/*.py.",
        "ÄÃ¡nh giÃ¡/XAI: src/utils/trainer.py, script/explain/*.py vÃ  src/evaluate/significance_tests.py.",
        "Káº¿t quáº£ lÆ°u trong checkpoint_ham10000, checkpoint_bcn20000 vÃ  results/significance_tests_auc.csv.",
    ]
    for item in sources:
        bullet(doc, item)

    heading(doc, "2. CÃ¡c Ä‘iá»ƒm pháº£i sá»­a ngay", 1)
    add_table(
        doc,
        ["Ná»™i dung slide/script/paper", "Káº¿t quáº£ Ä‘á»‘i chiáº¿u", "CÃ¡ch trÃ¬nh bÃ y nÃªn dÃ¹ng"],
        [
            (
                "PhÃ¢n loáº¡i 7-8 lá»›p bá»‡nh lÃ½.",
                "Model cÃ³ num_classes=1, dÃ¹ng sigmoid/BCE; HAM gá»™p mel, bcc, akiec thÃ nh Ã¡c tÃ­nh; BCN gÃ¡n Ã¡c tÃ­nh khi diagnosis_1 chá»©a 'malig'.",
                "BÃ i toÃ¡n hiá»‡n triá»ƒn khai lÃ  phÃ¢n loáº¡i nhá»‹ phÃ¢n: Ã¡c tÃ­nh so vá»›i khÃ´ng Ã¡c tÃ­nh."
            ),
            (
                "Patient-level split; chá»‘ng leakage tuyá»‡t Ä‘á»‘i.",
                "Cáº£ hai pipeline thá»±c táº¿ group báº±ng lesion_id. Code kiá»ƒm tra khÃ´ng trÃ¹ng lesion_id giá»¯a development/test vÃ  giá»¯a train/validation fold.",
                "Giao thá»©c lesion-level group splitting nháº±m háº¡n cháº¿ overlap tá»•n thÆ°Æ¡ng; khÃ´ng gá»i lÃ  patient-level náº¿u chÆ°a cÃ³ patient_id."
            ),
            (
                "ÄÃ£ loáº¡i biáº¿n sau cháº©n Ä‘oÃ¡n, vÃ­ dá»¥ káº¿t quáº£ sinh thiáº¿t.",
                "Paper nÃ³i Ä‘Ã£ loáº¡i histopathological confirmation/biopsy outcomes; nhÆ°ng BCN loader hiá»‡n dÃ¹ng diagnosis_confirm_type lÃ m metadata trong cÃ¡c mode cÃ³ metadata. ÄÃ¢y lÃ  mÃ¢u thuáº«n nghiÃªm trá»ng vá»›i paper.",
                "TrÃ¬nh bÃ y Ä‘Ã¢y lÃ  lá»—i cáº§n kháº¯c phá»¥c á»Ÿ báº£n code hiá»‡n táº¡i; loáº¡i diagnosis_confirm_type, huáº¥n luyá»‡n láº¡i vÃ  cáº­p nháº­t káº¿t quáº£ BCN náº¿u dÃ¹ng Ä‘á»ƒ báº£o vá»‡."
            ),
            (
                "So sÃ¡nh Ä‘áº§y Ä‘á»§ Image-only, Concatenation, FiLM, Gating.",
                "Paper Table 1 cÃ³ Ä‘á»§ 4 strategy trÃªn 4 backbone. Workspace hiá»‡n tháº¥y diag1=image-only, full/full_weighted=FiLM vÃ  late_fusion=gating; chÆ°a tháº¥y module concatenation thuáº§n.",
                "CÃ³ thá»ƒ nÃ³i Ä‘Ã¢y lÃ  káº¿t quáº£ cÃ´ng bá»‘ trong paper; náº¿u há»™i Ä‘á»“ng yÃªu cáº§u tÃ¡i láº­p tá»« code bÃ n giao, cáº§n bá»• sung code/checkpoint cá»§a concatenation vÃ  cÃ¡c backbone cÃ²n thiáº¿u."
            ),
            (
                "Káº¿t quáº£ ConvNeXt vÃ  ViT trÃªn slide.",
                "CÃ¡c sá»‘ Ä‘Æ°á»£c há»— trá»£ bá»Ÿi Table 2 cá»§a paper: HAM ConvNeXt+FiLM AUC 94,02% +/- 0,11; HAM ViT+FiLM 87,35% +/- 0,28; BCN ViT+FiLM 84,88% +/- 0,35. Workspace khÃ´ng chá»©a checkpoint/log tÆ°Æ¡ng á»©ng Ä‘á»ƒ kiá»ƒm tra láº¡i.",
                "Giá»¯ Ä‘Æ°á»£c trÃªn slide náº¿u ghi nguá»“n lÃ  Table 2 cá»§a paper; chuáº©n bá»‹ tráº£ lá»i ráº±ng artefact tÃ¡i láº­p Ä‘ang bÃ n giao trong workspace chÆ°a Ä‘áº§y Ä‘á»§ cho táº¥t cáº£ backbone."
            ),
            (
                "FiLM triá»‡t tiÃªu modality dominance vÃ  tÄƒng vá»t hiá»‡u nÄƒng.",
                "Paper cho tháº¥y lá»£i Ã­ch phá»¥ thuá»™c backbone: HAM ConvNeXt tÄƒng AUC tá»« 92,56 lÃªn 94,02, nhÆ°ng HAM EffNet chá»‰ 85,37 lÃªn 85,44. Trong log EffNet kiá»ƒm tra Ä‘Æ°á»£c, chÃªnh lá»‡ch khÃ´ng cÃ³ Ã½ nghÄ©a thá»‘ng kÃª (HAM p=0,853; BCN p=0,892).",
                "NÃ³i FiLM cho cáº£i thiá»‡n á»Ÿ má»™t sá»‘ backbone vÃ  duy trÃ¬ hiá»‡u nÄƒng cáº¡nh tranh á»Ÿ cÃ¡c backbone khÃ¡c; khÃ´ng nÃ³i triá»‡t tiÃªu hoÃ n toÃ n hay cáº£i thiá»‡n phá»• quÃ¡t."
            ),
            (
                "CÃ¹ng quy trÃ¬nh DullRazor, Gray-World, CLAHE cho cáº£ hai táº­p.",
                "Script tiá»n xá»­ lÃ½ HAM cÃ³ hair removal, illumination correction, Gray-World, CLAHE, smoothing; script BCN Ä‘ang dÃ¹ng hair removal nháº¹ vÃ  resize.",
                "MÃ´ táº£ preprocessing riÃªng cho tá»«ng dataset hoáº·c thá»‘ng nháº¥t pipeline rá»“i cháº¡y láº¡i."
            ),
            (
                "BCN20000 cÃ³ 19.424 áº£nh dÃ¹ng trong thÃ­ nghiá»‡m.",
                "Paper chá»‰ ghi xáº¥p xá»‰ 19.000; slide ghi 19.424; file metadata nguá»“n trong repo cÃ³ 18.946 dÃ²ng vÃ  CSV train/val/test thá»±c táº¿ cÃ³ tá»•ng 17.639 dÃ²ng.",
                "DÃ¹ng con sá»‘ truy váº¿t Ä‘Æ°á»£c: nguá»“n trong workspace 18.946; táº­p thá»±c nghiá»‡m hiá»‡n cÃ³ 17.639 sau chia/lá»c, náº¿u Ä‘Ãºng protocol cÃ´ng bá»‘."
            ),
        ],
        widths=[4.7, 7.2, 5.0],
    )

    heading(doc, "3. Pipeline thá»±c táº¿ theo mÃ£ nguá»“n", 1)
    heading(doc, "3.1. BÃ i toÃ¡n vÃ  nhÃ£n", 2)
    bullet(doc, "Äáº§u ra mÃ´ hÃ¬nh lÃ  má»™t logit; xÃ¡c suáº¥t Ã¡c tÃ­nh Ä‘Æ°á»£c láº¥y báº±ng sigmoid; threshold test máº·c Ä‘á»‹nh lÃ  0,5.")
    bullet(doc, "HAM10000: nhÃ£n Ã¡c tÃ­nh = {mel, bcc, akiec}; nhÃ£n khÃ´ng Ã¡c tÃ­nh = {nv, bkl, vasc, df}.")
    bullet(doc, "BCN: nhÃ£n Ã¡c tÃ­nh = diagnosis_1 chá»©a 'malig'; do logic hiá»‡n táº¡i, Indeterminate Ä‘Æ°á»£c Ä‘Æ°a vá» nhÃ£n 0. Quyáº¿t Ä‘á»‹nh nÃ y cáº§n Ä‘Æ°á»£c biá»‡n minh lÃ¢m sÃ ng hoáº·c sá»­a trÆ°á»›c khi cÃ´ng bá»‘.")

    heading(doc, "3.2. Dá»¯ liá»‡u thá»±c sá»± Ä‘i qua pipeline", 2)
    add_table(
        doc,
        ["Dataset/split", "Sá»‘ áº£nh", "Sá»‘ lesion_id", "Ãc tÃ­nh", "KhÃ´ng Ã¡c tÃ­nh", "Thiáº¿u tuá»•i"],
        [
            ("HAM train", "7.994", "5.976", "1.542", "6.452", "50"),
            ("HAM validation", "1.017", "747", "210", "807", "4"),
            ("HAM hold-out test", "1.004", "747", "202", "802", "3"),
            ("BCN train", "14.260", "4.023", "7.249", "7.011*", "111"),
            ("BCN validation", "1.704", "503", "799", "905*", "7"),
            ("BCN hold-out test", "1.675", "503", "823", "852*", "1"),
        ],
        widths=[3.5, 2.3, 2.4, 2.4, 3.0, 2.2],
    )
    p = doc.add_paragraph()
    p.add_run("* BCN 'khÃ´ng Ã¡c tÃ­nh' theo code hiá»‡n táº¡i gá»“m cáº£ cÃ¡c máº«u Indeterminate: ").bold = True
    p.add_run("875 train, 104 validation vÃ  109 test.")
    bullet(doc, "Kiá»ƒm tra trÃªn cÃ¡c CSV hiá»‡n cÃ³ cho tháº¥y overlap lesion_id giá»¯a train, validation vÃ  test báº±ng 0 á»Ÿ cáº£ HAM vÃ  BCN.")
    bullet(doc, "Trong script huáº¥n luyá»‡n, train vÃ  validation ban Ä‘áº§u Ä‘Æ°á»£c gá»™p thÃ nh development set, sau Ä‘Ã³ chia láº¡i báº±ng 5-fold StratifiedGroupKFold; hold-out test Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ sau má»—i fold.")

    heading(doc, "3.3. Metadata, fusion vÃ  loss", 2)
    add_table(
        doc,
        ["Háº¡ng má»¥c", "HAM10000", "BCN20000 / nháº­n xÃ©t"],
        [
            ("Metadata dÃ¹ng trong model", "age, localization, sex", "age_approx, anatom_site_general, anatom_site_special, diagnosis_confirm_type, sex"),
            ("Nguy cÆ¡ leakage", "dx vÃ  dx_type Ä‘Æ°á»£c drop trÆ°á»›c model.", "diagnosis_confirm_type váº«n lÃ  input: cáº§n loáº¡i bá» vÃ  train láº¡i."),
            ("Image-only", "Mode diag1", "Mode diag1"),
            ("FiLM", "Mode full; full_weighted dÃ¹ng meta_weight=2,0.", "Mode full; full_weighted dÃ¹ng meta_weight=2,0."),
            ("Gating", "Mode late_fusion (DualEmbeddingFusion).", "Mode late_fusion (DualEmbeddingFusion)."),
            ("Concatenation", "KhÃ´ng tháº¥y cÃ i Ä‘áº·t trong code hiá»‡n cÃ³.", "KhÃ´ng tháº¥y cÃ i Ä‘áº·t trong code hiá»‡n cÃ³."),
            ("Loss", "FocalLossBCE(alpha=0,75, gamma=2,0)", "FocalLossBCE(alpha=0,75, gamma=2,0)"),
        ],
        widths=[3.5, 6.0, 7.5],
    )
    bullet(doc, "FiLM Ä‘Æ°á»£c Ã¡p dá»¥ng lÃªn vector áº£nh sau backbone/global pooling, trÆ°á»›c classification head: F' = (1 + lambda*gamma) x F_img + lambda*beta.")
    bullet(doc, "lambda trong code lÃ  meta_weight: báº±ng 1,0 á»Ÿ mode full vÃ  2,0 á»Ÿ mode full_weighted; khÃ´ng pháº£i tham sá»‘ tá»± há»c.")
    bullet(doc, "Paper cÃ´ng bá»‘ ResNet50, EfficientNet-B4, ConvNeXt vÃ  ViT; workspace hiá»‡n cÃ³ code backbone EfficientNet-B4/ResNet50 vÃ  log CV kiá»ƒm tra trá»±c tiáº¿p cho EfficientNet-B4.")
    bullet(doc, "Hyperparameter chÆ°a khá»›p: paper ghi batch size 32 vÃ  learning rate HAM cÃ³ thá»ƒ lÃ  8 x 10^-5; train_ham.py/train_bcn.py hiá»‡n Ä‘áº·t batch size 16 vÃ  base learning rate 1 x 10^-4.")

    heading(doc, "4. Káº¿t quáº£ cÃ´ng bá»‘ trong bÃ i bÃ¡o CITA", 1)
    p = doc.add_paragraph(
        "CÃ¡c sá»‘ liá»‡u dÆ°á»›i Ä‘Ã¢y Ä‘Æ°á»£c chÃ©p láº¡i tá»« bÃ i bÃ¡o CITA_2026_paper_357.pdf. ÄÃ¢y lÃ  cÄƒn cá»© Ä‘á»ƒ "
        "giáº£i thÃ­ch báº£ng káº¿t quáº£ trong slide. Khi tráº£ lá»i há»™i Ä‘á»“ng, nÃªn gá»i Ä‘Ãºng lÃ  káº¿t quáº£ Ä‘Æ°á»£c bÃ¡o cÃ¡o "
        "trong paper; kháº£ nÄƒng tÃ¡i láº­p báº±ng workspace hiá»‡n táº¡i Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ riÃªng á»Ÿ má»¥c 5."
    )
    heading(doc, "4.1. Table 1 cá»§a paper: so sÃ¡nh bá»‘n chiáº¿n lÆ°á»£c", 2)
    add_table(
        doc,
        ["Dataset", "Architecture / Strategy", "AUC", "Accuracy", "F1-Score", "Specificity"],
        [
            ("HAM", "ResNet50 / S1 Image-only", "87,89%", "75,00%", "58,65%", "88,12%"),
            ("HAM", "ResNet50 / S2 Concatenation", "87,72%", "73,01%", "57,59%", "91,09%"),
            ("HAM", "ResNet50 / S3 FiLM", "88,52%", "71,61%", "56,88%", "93,07%"),
            ("HAM", "ResNet50 / S4 Gating", "85,51%", "72,31%", "56,43%", "89,11%"),
            ("HAM", "EffNet-B4 / S1 Image-only", "90,81%", "54,68%", "46,91%", "43,39%"),
            ("HAM", "EffNet-B4 / S2 Concatenation", "89,70%", "82,87%", "63,40%", "85,16%"),
            ("HAM", "EffNet-B4 / S3 FiLM", "90,02%", "81,18%", "63,44%", "81,17%"),
            ("HAM", "EffNet-B4 / S4 Gating", "90,34%", "61,16%", "50,51%", "51,75%"),
            ("HAM", "ConvNeXt / S1 Image-only", "91,95%", "69,12%", "55,71%", "62,22%"),
            ("HAM", "ConvNeXt / S2 Concatenation", "93,01%", "78,29%", "63,42%", "74,44%"),
            ("HAM", "ConvNeXt / S3 FiLM", "91,20%", "87,05%", "66,49%", "92,89%"),
            ("HAM", "ConvNeXt / S4 Gating", "91,96%", "86,06%", "67,44%", "89,65%"),
            ("HAM", "ViT / S1 Image-only", "85,72%", "77,45%", "79,14%", "80,91%"),
            ("HAM", "ViT / S2 Concatenation", "86,15%", "77,92%", "77,85%", "73,88%"),
            ("HAM", "ViT / S3 FiLM", "87,35%", "78,52%", "79,95%", "81,74%"),
            ("HAM", "ViT / S4 Gating", "82,95%", "75,18%", "75,40%", "76,55%"),
            ("BCN", "ResNet50 / S1 Image-only", "75,83%", "67,76%", "68,49%", "64,32%"),
            ("BCN", "ResNet50 / S2 Concatenation", "76,52%", "68,48%", "69,30%", "64,67%"),
            ("BCN", "ResNet50 / S3 FiLM", "76,41%", "68,30%", "68,93%", "65,14%"),
            ("BCN", "ResNet50 / S4 Gating", "72,62%", "65,19%", "66,78%", "59,39%"),
            ("BCN", "EffNet-B4 / S1 Image-only", "81,71%", "68,97%", "61,31%", "93,54%"),
            ("BCN", "EffNet-B4 / S2 Concatenation", "82,07%", "74,01%", "74,16%", "77,39%"),
            ("BCN", "EffNet-B4 / S3 FiLM", "82,00%", "72,92%", "74,24%", "71,47%"),
            ("BCN", "EffNet-B4 / S4 Gating", "81,07%", "73,12%", "73,83%", "74,16%"),
            ("BCN", "ConvNeXt / S1 Image-only", "81,18%", "73,43%", "73,75%", "71,01%"),
            ("BCN", "ConvNeXt / S2 Concatenation", "80,97%", "74,15%", "73,02%", "77,00%"),
            ("BCN", "ConvNeXt / S3 FiLM", "81,55%", "74,09%", "73,11%", "76,41%"),
            ("BCN", "ConvNeXt / S4 Gating", "82,02%", "73,55%", "73,71%", "71,71%"),
            ("BCN", "ViT / S1 Image-only", "84,24%", "76,18%", "77,81%", "79,47%"),
            ("BCN", "ViT / S2 Concatenation", "84,31%", "76,25%", "76,18%", "72,30%"),
            ("BCN", "ViT / S3 FiLM", "84,88%", "76,63%", "78,21%", "79,83%"),
            ("BCN", "ViT / S4 Gating", "81,24%", "73,73%", "73,72%", "74,97%"),
        ],
        widths=[1.8, 5.5, 2.4, 2.5, 2.5, 2.8],
    )
    heading(doc, "4.2. Table 2 cá»§a paper: 5-fold Strategy 1 so vá»›i Strategy 3", 2)
    add_table(
        doc,
        ["Dataset", "Architecture / Strategy", "AUC mean +/- SD", "Accuracy mean +/- SD", "F1 mean +/- SD", "Specificity mean +/- SD"],
        [
            ("HAM", "ResNet50 / S1", "87,49 +/- 0,08%", "72,85 +/- 1,21%", "57,46 +/- 0,85%", "68,25 +/- 1,78%"),
            ("HAM", "ResNet50 / S3 FiLM", "87,57 +/- 0,54%", "73,61 +/- 2,88%", "58,42 +/- 2,04%", "69,73 +/- 4,12%"),
            ("HAM", "EffNet-B4 / S1", "85,37 +/- 0,31%", "83,07 +/- 0,35%", "46,20 +/- 6,05%", "94,01 +/- 1,92%"),
            ("HAM", "EffNet-B4 / S3 FiLM", "85,44 +/- 1,01%", "84,77 +/- 0,65%", "46,32 +/- 4,86%", "94,26 +/- 2,41%"),
            ("HAM", "ConvNeXt / S1", "92,56 +/- 0,78%", "87,41 +/- 0,59%", "71,18 +/- 1,72%", "89,88 +/- 2,51%"),
            ("HAM", "ConvNeXt / S3 FiLM", "94,02 +/- 0,11%", "88,59 +/- 0,94%", "73,64 +/- 1,14%", "90,95 +/- 2,26%"),
            ("HAM", "ViT / S1", "85,72 +/- 0,42%", "77,45 +/- 0,55%", "79,14 +/- 0,62%", "80,91 +/- 1,05%"),
            ("HAM", "ViT / S3 FiLM", "87,35 +/- 0,28%", "78,52 +/- 0,42%", "79,95 +/- 0,48%", "81,74 +/- 0,92%"),
            ("BCN", "ResNet50 / S1", "76,41 +/- 0,39%", "68,70 +/- 0,33%", "68,07 +/- 0,95%", "67,97 +/- 3,12%"),
            ("BCN", "ResNet50 / S3 FiLM", "76,80 +/- 0,16%", "68,76 +/- 0,58%", "68,81 +/- 1,42%", "70,26 +/- 3,71%"),
            ("BCN", "EffNet-B4 / S1", "79,31 +/- 0,82%", "71,76 +/- 0,57%", "70,51 +/- 0,82%", "74,67 +/- 3,06%"),
            ("BCN", "EffNet-B4 / S3 FiLM", "79,65 +/- 0,83%", "71,98 +/- 0,95%", "71,09 +/- 1,03%", "73,71 +/- 4,49%"),
            ("BCN", "ConvNeXt / S1", "85,38 +/- 1,05%", "78,58 +/- 1,43%", "77,98 +/- 2,27%", "79,60 +/- 3,96%"),
            ("BCN", "ConvNeXt / S3 FiLM", "85,84 +/- 1,26%", "77,27 +/- 1,36%", "76,86 +/- 1,27%", "77,68 +/- 3,33%"),
            ("BCN", "ViT / S1", "84,24 +/- 0,58%", "76,18 +/- 0,65%", "77,81 +/- 0,82%", "79,47 +/- 1,25%"),
            ("BCN", "ViT / S3 FiLM", "84,88 +/- 0,35%", "76,63 +/- 0,52%", "78,21 +/- 0,64%", "79,83 +/- 1,10%"),
        ],
        widths=[1.8, 4.3, 3.0, 3.3, 3.2, 3.4],
    )
    heading(doc, "4.3. Nhá»¯ng chá»— cáº§n kiá»ƒm tra ngay trong paper", 2)
    bullet(doc, "Table 2 ghi BCN ConvNeXt Strategy 3 AUC = 85,84 +/- 1,26%, nhÆ°ng Ä‘oáº¡n vÄƒn á»Ÿ má»¥c 3.3 ghi 85,04 +/- 1,26%. Cáº§n sá»­a paper/slide/script dÃ¹ng thá»‘ng nháº¥t giÃ¡ trá»‹ tá»« báº£ng hoáº·c xÃ¡c minh láº¡i log gá»‘c.")
    bullet(doc, "Paper ghi dÃ¹ng strictly pre-diagnostic metadata vÃ  Ä‘Ã£ loáº¡i diagnostic confirmation, nhÆ°ng code BCN hiá»‡n váº«n dÃ¹ng diagnosis_confirm_type. ÄÃ¢y lÃ  mÃ¢u thuáº«n cÃ³ thá»ƒ áº£nh hÆ°á»Ÿng trá»±c tiáº¿p tÃ­nh há»£p lá»‡ cá»§a káº¿t quáº£ BCN.")
    bullet(doc, "Paper ghi protocol patient-level hoáº·c patient/lesion-level khÃ´ng nháº¥t quÃ¡n; code vÃ  CSV kiá»ƒm tra Ä‘Æ°á»£c lÃ  lesion_id. NÃªn gá»i chÃ­nh xÃ¡c lÃ  lesion-level group split trá»« khi cÃ³ patient_id thá»±c sá»±.")
    bullet(doc, "Paper ghi batch size 32 vÃ  learning rate HAM 8 x 10^-5 trong khi script train hiá»‡n cÃ³ batch size 16 vÃ  learning rate 1 x 10^-4. Cáº§n xÃ¡c nháº­n phiÃªn báº£n code táº¡o ra báº£ng paper.")

    heading(doc, "5. Káº¿t quáº£ cÃ³ thá»ƒ kiá»ƒm tra trá»±c tiáº¿p tá»« workspace hiá»‡n cÃ³", 1)
    p = doc.add_paragraph(
        "Báº£ng sau láº¥y tá»« cÃ¡c file test_metrics cá»§a nÄƒm model huáº¥n luyá»‡n theo tá»«ng fold. "
        "VÃ¬ cÃ¹ng hold-out test Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ bá»Ÿi nÄƒm model fold, nÃªn nÃªn mÃ´ táº£ lÃ  trung bÃ¬nh hiá»‡u nÄƒng "
        "cá»§a nÄƒm model trÃªn hold-out test, khÃ´ng diá»…n Ä‘áº¡t mÆ¡ há»“ lÃ  metric validation cá»§a 5 folds."
    )
    add_table(
        doc,
        ["Dataset", "Mode/code mapping", "AUC mean +/- SD", "Specificity mean +/- SD", "F1 mean +/- SD"],
        [
            ("HAM10000", "diag1 / image-only", "85,37% +/- 0,31", "94,71% +/- 1,92", "46,20% +/- 6,05"),
            ("HAM10000", "full / FiLM", "85,43% +/- 0,52", "93,24% +/- 3,33", "47,70% +/- 6,33"),
            ("HAM10000", "full_weighted / FiLM lambda=2", "85,44% +/- 1,01", "94,16% +/- 2,41", "46,32% +/- 4,86"),
            ("HAM10000", "late_fusion / gating", "82,54% +/- 1,04", "96,03% +/- 0,62", "33,81% +/- 3,95"),
            ("BCN20000", "diag1 / image-only", "79,31% +/- 0,82", "74,67% +/- 3,06", "70,51% +/- 0,82"),
            ("BCN20000", "full / FiLM", "79,31% +/- 0,26", "72,75% +/- 4,26", "70,81% +/- 1,44"),
            ("BCN20000", "full_weighted / FiLM lambda=2", "79,25% +/- 0,83", "73,71% +/- 4,49", "71,09% +/- 1,03"),
            ("BCN20000", "late_fusion / gating", "79,22% +/- 0,63", "75,45% +/- 3,66", "70,55% +/- 1,78"),
        ],
        widths=[2.7, 5.0, 3.4, 3.8, 3.4],
    )
    heading(doc, "5.1. Kiá»ƒm Ä‘á»‹nh Ã½ nghÄ©a AUC Ä‘Ã£ cÃ³ trong repo", 2)
    add_table(
        doc,
        ["So sÃ¡nh", "ChÃªnh lá»‡ch AUC", "p-value (paired t-test)", "Diá»…n giáº£i Ä‘Ãºng"],
        [
            ("HAM: FiLM weighted - image-only", "+0,075 Ä‘iá»ƒm pháº§n trÄƒm", "0,8531", "KhÃ´ng cÃ³ Ã½ nghÄ©a thá»‘ng kÃª."),
            ("HAM: gating - image-only", "-2,822 Ä‘iá»ƒm pháº§n trÄƒm", "0,0077", "Gating tháº¥p hÆ¡n image-only trong log hiá»‡n cÃ³."),
            ("BCN: FiLM weighted - image-only", "-0,062 Ä‘iá»ƒm pháº§n trÄƒm", "0,8920", "KhÃ´ng cÃ³ Ã½ nghÄ©a thá»‘ng kÃª."),
            ("BCN: gating - image-only", "-0,089 Ä‘iá»ƒm pháº§n trÄƒm", "0,6205", "KhÃ´ng cÃ³ Ã½ nghÄ©a thá»‘ng kÃª."),
        ],
        widths=[5.2, 3.7, 3.4, 5.6],
    )
    p = doc.add_paragraph()
    p.add_run("CÃ¢u káº¿t quáº£ nÃªn nÃ³i: ").bold = True
    p.add_run(
        "â€œTrong cáº¥u hÃ¬nh EfficientNet-B4 Ä‘Æ°á»£c tÃ¡i láº­p tá»« mÃ£ nguá»“n, metadata/FiLM chÆ°a cáº£i thiá»‡n AUC "
        "cÃ³ Ã½ nghÄ©a thá»‘ng kÃª so vá»›i image-only. Káº¿t quáº£ cho tháº¥y nhu cáº§u kiá»ƒm soÃ¡t leakage, chuáº©n hÃ³a "
        "metadata vÃ  má»Ÿ rá»™ng Ä‘á»‘i chá»©ng trÆ°á»›c khi káº¿t luáº­n vá» Æ°u tháº¿ cá»§a cÆ¡ cháº¿ fusion.â€"
    )
    p = doc.add_paragraph()
    p.add_run("Lá»‡ch giá»¯a paper vÃ  workspace cáº§n nÃ³i rÃµ: ").bold = True
    p.add_run(
        "Table 2 cá»§a paper bÃ¡o cÃ¡o BCN EffNet-B4 Strategy 3 AUC 79,65% +/- 0,83, trong khi file "
        "checkpoint hiá»‡n Ä‘á»‘i chiáº¿u cho mode `full_weighted` cho AUC 79,25% +/- 0,83. Hai cáº¥u hÃ¬nh "
        "khÃ´ng nÃªn Ä‘Æ°á»£c coi lÃ  cÃ¹ng má»™t láº§n cháº¡y náº¿u chÆ°a truy xuáº¥t Ä‘Ãºng log/config táº¡o báº£ng paper."
    )

    heading(doc, "6. Sá»­a slide vÃ  script theo tá»«ng vá»‹ trÃ­", 1)
    add_table(
        doc,
        ["Vá»‹ trÃ­", "Váº¥n Ä‘á»", "Ná»™i dung thay tháº¿ Ä‘á» xuáº¥t"],
        [
            (
                "Slide/script 3",
                "NÃ³i SOTA AUC >95% do leakage nhÆ° má»™t káº¿t luáº­n phá»• quÃ¡t.",
                "Äá»•i thÃ nh: 'Chia theo áº£nh cÃ³ thá»ƒ gÃ¢y leakage khi cÃ¹ng tá»•n thÆ°Æ¡ng xuáº¥t hiá»‡n á»Ÿ nhiá»u áº£nh; do Ä‘Ã³ nghiÃªn cá»©u dÃ¹ng group split theo lesion_id.' Chá»‰ giá»¯ sá»‘ AUC náº¿u cÃ³ trÃ­ch dáº«n."
            ),
            (
                "Slide/script 4",
                "Tá»« tuyá»‡t Ä‘á»‘i: chá»‘ng leakage tuyá»‡t Ä‘á»‘i, chá»©ng minh cÆ¡ sá»Ÿ y khoa.",
                "Äá»•i thÃ nh: 'Háº¡n cháº¿ overlap tá»•n thÆ°Æ¡ng báº±ng lesion-level splitting vÃ  dÃ¹ng XAI Ä‘á»ƒ kháº£o sÃ¡t Ä‘á»‹nh tÃ­nh vÃ¹ng chÃº Ã½/Ä‘Ã³ng gÃ³p metadata.'"
            ),
            (
                "Slide/script 5",
                "Sai task 7-8 lá»›p; sai/khÃ´ng rÃµ quy mÃ´ BCN.",
                "Äá»•i thÃ nh binary malignant vs non-malignant. Ghi HAM=10.015 áº£nh; BCN: nguá»“n repo=18.946, táº­p CSV thÃ­ nghiá»‡m=17.639 áº£nh."
            ),
            (
                "Slide/script 6",
                "Gá»i patient-level vÃ  nÃ³i Ä‘Ã£ loáº¡i biáº¿n háº­u cháº©n Ä‘oÃ¡n.",
                "Äá»•i thÃ nh lesion-level. ThÃªm giá»›i háº¡n: BCN cáº§n loáº¡i diagnosis_confirm_type khá»i input vÃ  cháº¡y láº¡i trÆ°á»›c khi bÃ¡o cÃ¡o káº¿t quáº£ metadata."
            ),
            (
                "Slide/script 7",
                "Bá»‘n strategy cÃ³ trong paper Table 1, nhÆ°ng code workspace thiáº¿u Concatenation.",
                "Giá»¯ bá»‘n strategy vá»›i chÃº thÃ­ch 'káº¿t quáº£ cÃ´ng bá»‘ trong paper'; chuáº©n bá»‹ bá»• sung mÃ£/checkpoint Strategy 2 náº¿u há»™i Ä‘á»“ng yÃªu cáº§u tÃ¡i láº­p."
            ),
            (
                "Slide/script 9",
                "KhÃ´ng giáº£i thÃ­ch lambda trong cÃ´ng thá»©c FiLM.",
                "Ghi rÃµ: modulation sau global pooling; lambda=meta_weight Ä‘áº·t trÆ°á»›c (1 hoáº·c 2), gamma/beta sinh tá»« metadata."
            ),
            (
                "Slide/script 10",
                "Paper cÃ´ng bá»‘ bá»‘n backbone nhÆ°ng workspace hiá»‡n khÃ´ng bÃ n giao Ä‘á»§ artefact cho ConvNeXt/ViT.",
                "NÃªu bá»‘n backbone lÃ  pháº¡m vi thá»±c nghiá»‡m cá»§a paper; khÃ´ng nÃ³i toÃ n bá»™ checkpoint Ä‘Ã£ sáºµn sÃ ng trong mÃ£ nguá»“n náº¿u chÆ°a bá»• sung."
            ),
            (
                "Slide/script 11",
                "Slide trÃ­ch má»™t pháº§n Table 2 cá»§a paper nhÆ°ng káº¿t luáº­n Ä‘ang quÃ¡ tuyá»‡t Ä‘á»‘i.",
                "Ghi nguá»“n Table 2 cá»§a paper. NÃ³i FiLM cáº£i thiá»‡n rÃµ á»Ÿ HAM ConvNeXt/ViT vÃ  duy trÃ¬ cáº¡nh tranh á»Ÿ cÃ¡c cáº¥u hÃ¬nh khÃ¡c; khÃ´ng suy rá»™ng thÃ nh cáº£i thiá»‡n phá»• quÃ¡t hoáº·c hiá»‡u quáº£ lÃ¢m sÃ ng Ä‘Ã£ chá»©ng minh."
            ),
            (
                "Slide/script 12",
                "DÃ¹ng Grad-CAM/SHAP Ä‘á»ƒ 'kháº³ng Ä‘á»‹nh' cÆ¡ cháº¿ Ä‘Ãºng.",
                "Äá»•i thÃ nh: 'Grad-CAM vÃ  SHAP lÃ  phÃ¢n tÃ­ch há»— trá»£/kháº£o sÃ¡t; chÆ°a pháº£i xÃ¡c nháº­n lÃ¢m sÃ ng hay quan há»‡ nhÃ¢n quáº£.'"
            ),
            (
                "Slide/script 13-14",
                "TuyÃªn bá»‘ leakage-free cáº¥p bá»‡nh nhÃ¢n vÃ  triá»‡t tiÃªu modality dominance.",
                "Äá»•i thÃ nh Ä‘Ã³ng gÃ³p triá»ƒn khai: pipeline binary Ä‘a phÆ°Æ¡ng thá»©c, lesion-level split, FiLM/gating vÃ  phÃ¢n tÃ­ch giá»›i háº¡n cáº§n kháº¯c phá»¥c."
            ),
        ],
        widths=[3.0, 5.6, 8.7],
    )

    heading(doc, "7. CÃ¡c lá»—i ká»¹ thuáº­t cáº§n xá»­ lÃ½ trÆ°á»›c khi chá»‘t bÃ¡o cÃ¡o", 1)
    numbered(doc, "Loáº¡i `diagnosis_confirm_type` khá»i danh sÃ¡ch metadata cá»§a BCN, vÃ¬ Ä‘Ã¢y lÃ  thÃ´ng tin xÃ¡c nháº­n cháº©n Ä‘oÃ¡n; sau Ä‘Ã³ huáº¥n luyá»‡n láº¡i má»i mode cÃ³ metadata vÃ  cáº­p nháº­t báº£ng.")
    numbered(doc, "Quyáº¿t Ä‘á»‹nh rÃµ cÃ¡ch xá»­ lÃ½ nhÃ£n BCN `Indeterminate`: loáº¡i khá»i binary evaluation hoáº·c Ä‘á»‹nh nghÄ©a chÃ­nh thá»©c vÃ¬ sao xem lÃ  negative; sau Ä‘Ã³ cháº¡y láº¡i náº¿u thay Ä‘á»•i.")
    numbered(doc, "Bá»• sung artefact thá»±c nghiá»‡m cá»§a Strategy 2/ConvNeXt/ViT náº¿u cáº§n chá»©ng minh kháº£ nÄƒng tÃ¡i láº­p cÃ¡c káº¿t quáº£ Ä‘Ã£ cÃ´ng bá»‘ trong paper.")
    numbered(doc, "Sá»­a mÃ¢u thuáº«n ná»™i bá»™ paper á»Ÿ BCN ConvNeXt Strategy 3: Table 2 ghi AUC 85,84%, nhÆ°ng pháº§n so sÃ¡nh SOTA ghi 85,04%.")
    numbered(doc, "Äá»‘i chiáº¿u láº¡i hyperparameter paper vá»›i code bÃ n giao: batch size 32 so vá»›i 16 vÃ  learning rate HAM 8 x 10^-5 so vá»›i 1 x 10^-4.")
    numbered(doc, "Thá»‘ng nháº¥t preprocessing giá»¯a lá»i trÃ¬nh bÃ y vÃ  code: HAM hiá»‡n cÃ³ Gray-World/CLAHE, BCN hiá»‡n khÃ´ng cÃ³ cÃ¡c bÆ°á»›c nÃ y.")
    numbered(doc, "Sá»­a cÃ¡ch diá»…n Ä‘áº¡t CV/test: code táº¡o 5 model báº±ng CV trÃªn development set vÃ  cÃ¹ng Ä‘Ã¡nh giÃ¡ trÃªn hold-out test.")
    numbered(doc, "Sá»­a hoáº·c kiá»ƒm chá»©ng script SHAP/figure trÆ°á»›c khi trÃ¬nh chiáº¿u; SHAP hiá»‡n lÃ  phÃ¢n tÃ­ch há»— trá»£, khÃ´ng chá»©ng minh hiá»‡u nÄƒng hay tÃ­nh Ä‘Ãºng y khoa.")

    heading(doc, "8. CÃ¢u há»i há»™i Ä‘á»“ng cÃ³ kháº£ nÄƒng há»i cao vÃ  Ä‘Ã¡p Ã¡n Ä‘Ã£ hiá»‡u Ä‘Ã­nh", 1)
    add_qa(
        doc,
        "BÃ i toÃ¡n cá»§a nhÃ³m lÃ  binary hay multi-class?",
        "Theo mÃ£ nguá»“n hiá»‡n táº¡i, Ä‘Ã¢y lÃ  bÃ i toÃ¡n binary malignant/non-malignant. MÃ´ hÃ¬nh cÃ³ má»™t Ä‘áº§u ra logit, xÃ¡c suáº¥t láº¥y báº±ng sigmoid vÃ  test dÃ¹ng threshold 0,5.",
        "KhÃ´ng nÃ³i phÃ¢n loáº¡i 7-8 lá»›p náº¿u chÆ°a Ä‘á»•i model/loss/evaluation."
    )
    add_qa(
        doc,
        "NhÃ³m chá»‘ng leakage á»Ÿ cáº¥p bá»‡nh nhÃ¢n hay tá»•n thÆ°Æ¡ng?",
        "Trong dá»¯ liá»‡u vÃ  code hiá»‡n táº¡i, nhÃ³m dÃ¹ng `lesion_id` lÃ m group. NhÃ³m kiá»ƒm tra khÃ´ng trÃ¹ng lesion_id giá»¯a development set vÃ  hold-out test cÅ©ng nhÆ° giá»¯a train/validation trong má»—i fold.",
        "ÄÃ¢y lÃ  lesion-level protection; khÃ´ng Ä‘á»“ng nháº¥t vá»›i patient-level náº¿u thiáº¿u patient_id."
    )
    add_qa(
        doc,
        "Metadata gá»“m nhá»¯ng biáº¿n nÃ o?",
        "HAM sá»­ dá»¥ng age, sex vÃ  localization. BCN code hiá»‡n náº¡p age_approx, sex, anatom_site_general, anatom_site_special vÃ  diagnosis_confirm_type.",
        "Biáº¿n diagnosis_confirm_type cá»§a BCN cÃ³ nguy cÆ¡ háº­u cháº©n Ä‘oÃ¡n vÃ  cáº§n bá» trÆ°á»›c khi cÃ´ng bá»‘ káº¿t quáº£ metadata sáº¡ch."
    )
    add_qa(
        doc,
        "FiLM Ä‘Æ°á»£c Ä‘áº·t á»Ÿ Ä‘Ã¢u vÃ  lambda lÃ  gÃ¬?",
        "FiLM Ä‘iá»u biáº¿n vector feature áº£nh sau backbone/global pooling vÃ  trÆ°á»›c classifier. CÃ´ng thá»©c code lÃ  F'=(1+lambda*gamma)F+lambda*beta; lambda chÃ­nh lÃ  `meta_weight`, báº±ng 1 á»Ÿ `full` vÃ  2 á»Ÿ `full_weighted`."
    )
    add_qa(
        doc,
        "Bá»‘n chiáº¿n lÆ°á»£c fusion Ä‘Ã£ Ä‘Æ°á»£c triá»ƒn khai Ä‘áº§y Ä‘á»§ chÆ°a?",
        "BÃ i bÃ¡o cÃ´ng bá»‘ Table 1 so sÃ¡nh Image-only, Concatenation, FiLM vÃ  Gating trÃªn bá»‘n backbone. Trong workspace Ä‘ang Ä‘á»‘i chiáº¿u, em má»›i truy váº¿t trá»±c tiáº¿p Ä‘Æ°á»£c code/checkpoint cá»§a Image-only, FiLM/FiLM-weighted vÃ  Gating cho EfficientNet-B4; artefact cá»§a Concatenation vÃ  má»™t sá»‘ backbone cáº§n bá»• sung náº¿u nghiá»‡m thu yÃªu cáº§u cháº¡y láº¡i."
    )
    add_qa(
        doc,
        "VÃ¬ sao dÃ¹ng Focal BCE Loss?",
        "VÃ¬ bÃ i toÃ¡n code lÃ  binary vÃ  HAM máº¥t cÃ¢n báº±ng giá»¯a malignant vÃ  non-malignant. Cáº¥u hÃ¬nh hiá»‡n dÃ¹ng alpha=0,75 vÃ  gamma=2,0 Ä‘á»ƒ giáº£m áº£nh hÆ°á»Ÿng máº«u dá»… vÃ  táº­p trung hÆ¡n vÃ o lá»—i khÃ³.",
        "KhÃ´ng nÃ³i cÃ¡c tham sá»‘ Ä‘Ã£ tá»‘i Æ°u náº¿u chÆ°a cÃ³ ablation/tuning log."
    )
    add_qa(
        doc,
        "FiLM cÃ³ tá»‘t hÆ¡n image-only khÃ´ng?",
        "Theo Table 2 cá»§a paper, má»©c cáº£i thiá»‡n phá»¥ thuá»™c kiáº¿n trÃºc: HAM ConvNeXt tÄƒng tá»« 92,56% lÃªn 94,02% vÃ  ViT tÄƒng tá»« 85,72% lÃªn 87,35%; EffNet-B4 chá»‰ tÄƒng nháº¹ tá»« 85,37% lÃªn 85,44%. Vá»›i log EfficientNet-B4 hiá»‡n cÃ³ trong workspace, kiá»ƒm Ä‘á»‹nh bá»• sung chÆ°a cho tháº¥y khÃ¡c biá»‡t AUC cÃ³ Ã½ nghÄ©a. VÃ¬ váº­y nhÃ³m káº¿t luáº­n tháº­n trá»ng ráº±ng FiLM cÃ³ lá»£i á»Ÿ má»™t sá»‘ backbone, khÃ´ng pháº£i vÆ°á»£t trá»™i phá»• quÃ¡t."
    )
    add_qa(
        doc,
        "Specificity cao cÃ³ nghÄ©a lÃ  giáº£m sinh thiáº¿t khÃ´ng cáº§n thiáº¿t khÃ´ng?",
        "Specificity cao cho tháº¥y mÃ´ hÃ¬nh giáº£m false positive táº¡i threshold Ä‘ang dÃ¹ng, nÃªn cÃ³ tiá»m nÄƒng há»— trá»£ giáº£m chá»‰ Ä‘á»‹nh khÃ´ng cáº§n thiáº¿t. Tuy nhiÃªn chÆ°a thá»ƒ suy ra tÃ¡c Ä‘á»™ng lÃ¢m sÃ ng trá»±c tiáº¿p náº¿u chÆ°a cÃ³ Ä‘Ã¡nh giÃ¡ tiáº¿n cá»©u vá»›i bÃ¡c sÄ©."
    )
    add_qa(
        doc,
        "Grad-CAM vÃ  SHAP chá»©ng minh Ä‘Æ°á»£c gÃ¬?",
        "Grad-CAM giÃºp quan sÃ¡t vÃ¹ng áº£nh áº£nh hÆ°á»Ÿng tá»›i dá»± Ä‘oÃ¡n; SHAP kháº£o sÃ¡t Ä‘Ã³ng gÃ³p tÆ°Æ¡ng Ä‘á»‘i cá»§a feature metadata trong thiáº¿t láº­p phÃ¢n tÃ­ch. ChÃºng lÃ  cÃ´ng cá»¥ há»— trá»£ giáº£i thÃ­ch, khÃ´ng chá»©ng minh quan há»‡ nhÃ¢n quáº£ hay Ä‘á»™ tin cáº­y lÃ¢m sÃ ng."
    )
    add_qa(
        doc,
        "Táº¡i sao káº¿t quáº£ slide cÃ³ ConvNeXt/ViT nhÆ°ng mÃ£ nguá»“n chá»‰ thá»ƒ hiá»‡n EfficientNet-B4/ResNet50?",
        "ConvNeXt/ViT lÃ  káº¿t quáº£ Ä‘Æ°á»£c cÃ´ng bá»‘ trong Table 1 vÃ  Table 2 cá»§a paper CITA. Tuy nhiÃªn, workspace hiá»‡n táº¡i chÆ°a kÃ¨m Ä‘áº§y Ä‘á»§ checkpoint/log tÆ°Æ¡ng á»©ng Ä‘á»ƒ tÃ¡i cháº¡y toÃ n bá»™ báº£ng. NhÃ³m cáº§n phÃ¢n biá»‡t káº¿t quáº£ cÃ´ng bá»‘ vÃ  pháº¡m vi artefact hiá»‡n bÃ n giao, Ä‘á»“ng thá»i bá»• sung artefact náº¿u há»™i Ä‘á»“ng yÃªu cáº§u nghiá»‡m thu tÃ¡i láº­p."
    )

    heading(doc, "9. PhiÃªn báº£n lá»i nÃ³i ngáº¯n Ä‘á»ƒ thay trong script", 1)
    replacements = [
        (
            "Má»Ÿ Ä‘áº§u/má»¥c tiÃªu",
            "NhÃ³m xÃ¢y dá»±ng mÃ´ hÃ¬nh há»— trá»£ phÃ¢n loáº¡i nguy cÆ¡ Ã¡c tÃ­nh tá»« áº£nh dermoscopy káº¿t há»£p metadata cÃ³ cáº¥u trÃºc. Má»¥c tiÃªu lÃ  kháº£o sÃ¡t liá»‡u metadata vÃ  cÆ¡ cháº¿ fusion cÃ³ há»— trá»£ mÃ´ hÃ¬nh trong má»™t protocol háº¡n cháº¿ overlap tá»•n thÆ°Æ¡ng hay khÃ´ng."
        ),
        (
            "Dá»¯ liá»‡u",
            "TrÃªn HAM10000, chÃºng em sá»­ dá»¥ng Ä‘á»§ 10.015 áº£nh. Vá»›i BCN, file nguá»“n trong workspace cÃ³ 18.946 báº£n ghi vÃ  cÃ¡c split Ä‘Æ°á»£c Ä‘Æ°a vÃ o thÃ­ nghiá»‡m cÃ³ tá»•ng 17.639 áº£nh. NhÃ£n triá»ƒn khai lÃ  nhá»‹ phÃ¢n Ã¡c tÃ­nh/khÃ´ng Ã¡c tÃ­nh."
        ),
        (
            "Protocol",
            "ChÃºng em dÃ¹ng `lesion_id` Ä‘á»ƒ group, kiá»ƒm tra khÃ´ng cÃ³ lesion trÃ¹ng giá»¯a táº­p development vÃ  test, sau Ä‘Ã³ dÃ¹ng 5-fold Stratified Group CV trÃªn development set. Má»—i model fold Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ trÃªn cÃ¹ng hold-out test."
        ),
        (
            "FiLM",
            "FiLM nháº­n metadata vÃ  sinh gamma, beta Ä‘á»ƒ Ä‘iá»u biáº¿n feature áº£nh sau backbone. Mode weighted Ä‘áº·t há»‡ sá»‘ metadata báº±ng 2; Ä‘Ã¢y lÃ  cáº¥u hÃ¬nh Ä‘áº·t trÆ°á»›c chá»© khÃ´ng pháº£i trá»ng sá»‘ tá»± há»c."
        ),
        (
            "Káº¿t quáº£",
            "Theo Table 2 cá»§a bÃ i bÃ¡o, FiLM cho má»©c cáº£i thiá»‡n phá»¥ thuá»™c backbone. Ná»•i báº­t trÃªn HAM10000, ConvNeXt + FiLM Ä‘áº¡t AUC 94,02% +/- 0,11 so vá»›i 92,56% +/- 0,78 cá»§a image-only; ViT + FiLM Ä‘áº¡t 87,35% +/- 0,28. Vá»›i EfficientNet-B4, cáº£i thiá»‡n AUC nhá» hÆ¡n, nÃªn nhÃ³m khÃ´ng kháº³ng Ä‘á»‹nh FiLM luÃ´n vÆ°á»£t trá»™i."
        ),
        (
            "Giá»›i háº¡n",
            "Äáº·c biá»‡t, pipeline BCN hiá»‡n cÃ²n biáº¿n `diagnosis_confirm_type` cÃ³ nguy cÆ¡ leakage vÃ  cáº§n loáº¡i bá» rá»“i huáº¥n luyá»‡n láº¡i. VÃ¬ váº­y káº¿t quáº£ BCN Ä‘a phÆ°Æ¡ng thá»©c Ä‘Æ°á»£c xem lÃ  káº¿t quáº£ sÆ¡ bá»™, chÆ°a dÃ¹ng Ä‘á»ƒ kháº³ng Ä‘á»‹nh cáº£i thiá»‡n lÃ¢m sÃ ng."
        ),
    ]
    add_table(doc, ["Äoáº¡n", "CÃ¢u nÃ³i Ä‘á» xuáº¥t"], replacements, widths=[3.4, 13.8])

    heading(doc, "10. Checklist trÆ°á»›c khi báº£o vá»‡", 1)
    checks = [
        "Äá»•i toÃ n bá»™ cá»¥m '7-8 lá»›p' thÃ nh bÃ i toÃ¡n binary hoáº·c thay Ä‘á»•i code/thÃ­ nghiá»‡m tÆ°Æ¡ng á»©ng.",
        "Sá»­a má»i cá»¥m 'patient-level' thÃ nh 'lesion-level' theo code hiá»‡n táº¡i.",
        "KhÃ´ng dÃ¹ng káº¿t quáº£ BCN Ä‘a phÆ°Æ¡ng thá»©c nhÆ° báº±ng chá»©ng máº¡nh trÆ°á»›c khi loáº¡i diagnosis_confirm_type vÃ  cháº¡y láº¡i.",
        "Gáº¯n nhÃ£n báº£ng slide lÃ  trÃ­ch tá»« Table 2 cá»§a paper; cung cáº¥p Ä‘áº§y Ä‘á»§ log/checkpoint náº¿u há»™i Ä‘á»“ng yÃªu cáº§u tÃ¡i láº­p backbone khÃ¡c.",
        "Sá»­a lá»—i paper vá» BCN ConvNeXt Strategy 3: xÃ¡c minh 85,84% hay 85,04% lÃ  giÃ¡ trá»‹ Ä‘Ãºng.",
        "KhÃ´ng káº¿t luáº­n FiLM vÆ°á»£t trá»™i phá»• quÃ¡t hoáº·c triá»‡t tiÃªu modality dominance; paper cho tháº¥y hiá»‡u quáº£ phá»¥ thuá»™c backbone.",
        "LÃ m rÃµ cÃ¡c máº«u BCN Indeterminate Ä‘Æ°á»£c xá»­ lÃ½ nhÆ° tháº¿ nÃ o.",
        "Cáº­p nháº­t mÃ´ táº£ preprocessing riÃªng biá»‡t cho HAM vÃ  BCN.",
        "Chuáº©n bá»‹ nguá»“n trÃ­ch dáº«n cho phÃ¡t biá»ƒu vá» SOTA vÃ  tÃ¡c Ä‘á»™ng lÃ¢m sÃ ng.",
    ]
    for check in checks:
        bullet(doc, check)

    heading(doc, "11. Káº¿t luáº­n hiá»‡u Ä‘Ã­nh", 1)
    doc.add_paragraph(
        "TÃ i liá»‡u trá»ng tÃ¢m ban Ä‘áº§u Ä‘Ãºng á»Ÿ viá»‡c cáº£nh bÃ¡o mÃ¢u thuáº«n binary/multi-class, yÃªu cáº§u lÃ m rÃµ leakage, "
        "loss vÃ  báº±ng chá»©ng thá»±c nghiá»‡m. Paper CITA Ä‘Ã£ cung cáº¥p báº£ng káº¿t quáº£ bá»‘n strategy/bá»‘n backbone nÃªn cÃ¡c "
        "sá»‘ ConvNeXt vÃ  ViT trÃªn slide cÃ³ nguá»“n cÃ´ng bá»‘. DÃ¹ váº­y, nhÃ³m pháº£i sá»­a cÃ¡ch diá»…n Ä‘áº¡t thÃ nh bÃ i toÃ¡n binary, "
        "lÃ m rÃµ lesion-level protocol, xá»­ lÃ½ mÃ¢u thuáº«n `diagnosis_confirm_type` trong BCN, xÃ¡c minh sai lá»‡ch sá»‘ "
        "BCN ConvNeXt trong paper vÃ  phÃ¢n biá»‡t káº¿t quáº£ cÃ´ng bá»‘ vá»›i artefact tÃ¡i láº­p hiá»‡n Ä‘ang bÃ n giao. "
        "ÄÃ¢y lÃ  cÃ¡ch trÃ¬nh bÃ y trung thá»±c, ká»¹ thuáº­t vá»¯ng vÃ  an toÃ n nháº¥t trÆ°á»›c há»™i Ä‘á»“ng."
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_document()
