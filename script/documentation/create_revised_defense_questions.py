from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path("deliverables/Bo_cau_hoi_phan_bien_DA_CHINH_SUA.docx")


QA_ITEMS = [
    (
        "1. RÃ² rá»‰ dá»¯ liá»‡u á»Ÿ cáº¥p Ä‘á»™ bá»‡nh nhÃ¢n",
        "\"Trong slide 3 vÃ  slide 6, nhÃ³m tá»± tin kháº³ng Ä‘á»‹nh dÃ¹ng giao thá»©c chia dá»¯ liá»‡u theo lesion_id Ä‘á»ƒ chá»‘ng rÃ² rá»‰ dá»¯ liá»‡u (Data Leakage). Tuy nhiÃªn, vá» máº·t y khoa, má»™t bá»‡nh nhÃ¢n cÃ³ thá»ƒ cÃ³ nhiá»u tá»•n thÆ°Æ¡ng vá»›i cÃ¡c mÃ£ lesion_id khÃ¡c nhau trÃªn cÃ¹ng má»™t cÆ¡ thá»ƒ. Chia theo lesion_id hoÃ n toÃ n khÃ´ng ngÄƒn cháº·n Ä‘Æ°á»£c viá»‡c áº£nh cá»§a cÃ¹ng má»™t bá»‡nh nhÃ¢n xuáº¥t hiá»‡n á»Ÿ cáº£ táº­p Train vÃ  Test. Táº¡i sao nhÃ³m láº¡i bá» qua rá»§i ro Patient-level Leakage nÃ y?\"",
        [
            "NhÃ³m khÃ´ng phá»§ nháº­n rá»§i ro nÃ y. Giao thá»©c hiá»‡n táº¡i kiá»ƒm soÃ¡t rÃ² rá»‰ á»Ÿ cáº¥p Ä‘á»™ tá»•n thÆ°Æ¡ng: má»i áº£nh cÃ³ cÃ¹ng lesion_id Ä‘Æ°á»£c giá»¯ trong cÃ¹ng má»™t pháº§n dá»¯ liá»‡u. Äiá»u Ä‘Ã³ ngÄƒn má»™t tá»•n thÆ°Æ¡ng xuáº¥t hiá»‡n Ä‘á»“ng thá»i á»Ÿ táº­p há»c vÃ  táº­p Ä‘Ã¡nh giÃ¡, nhÆ°ng khÃ´ng tÆ°Æ¡ng Ä‘Æ°Æ¡ng vá»›i patient-level split.",
            "Trong cÃ¡c tá»‡p metadata Ä‘ang sá»­ dá»¥ng cho HAM10000 vÃ  BCN20000 cá»§a dá»± Ã¡n, nhÃ³m cÃ³ trÆ°á»ng lesion_id nhÆ°ng khÃ´ng cÃ³ patient_id. VÃ¬ váº­y, vá»›i dá»¯ liá»‡u hiá»‡n cÃ³, nhÃ³m chÆ°a thá»ƒ chá»©ng minh ráº±ng hai tá»•n thÆ°Æ¡ng khÃ¡c nhau cá»§a cÃ¹ng bá»‡nh nhÃ¢n khÃ´ng bá»‹ phÃ¢n bá»‘ sang hai táº­p khÃ¡c nhau.",
            "Káº¿t luáº­n Ä‘Ãºng lÃ : nhÃ³m Ä‘Ã£ giáº£m má»™t nguá»“n leakage quan trá»ng á»Ÿ cáº¥p lesion, nhÆ°ng patient-level leakage váº«n lÃ  giá»›i háº¡n dá»¯ liá»‡u cáº§n cÃ´ng bá»‘ minh báº¡ch. Náº¿u nháº­n Ä‘Æ°á»£c patient_id hoáº·c bá»™ dá»¯ liá»‡u cÃ³ Ä‘á»‹nh danh bá»‡nh nhÃ¢n, giao thá»©c Æ°u tiÃªn pháº£i lÃ  patient-level group split. Káº¿t quáº£ cross-dataset chá»‰ Ä‘Æ°á»£c dÃ¹ng nhÆ° Ä‘Ã¡nh giÃ¡ kháº£ nÄƒng tá»•ng quÃ¡t hÃ³a khi Ä‘Ã£ cháº¡y vÃ  bÃ¡o cÃ¡o riÃªng, khÃ´ng Ä‘Æ°á»£c dÃ¹ng Ä‘á»ƒ tuyÃªn bá»‘ loáº¡i bá» patient-level leakage.",
        ],
    ),
    (
        "2. FiLM phá»¥ thuá»™c vÃ o backbone",
        "\"HÃ£y nhÃ¬n vÃ o báº£ng káº¿t quáº£ trÃªn táº­p HAM10000: Khi cÃ¹ng tÃ­ch há»£p cÆ¡ cháº¿ FiLM (Strategy 3), ConvNeXt Ä‘áº¡t 94.02% AUC, nhÆ°ng EfficientNet-B4 chá»‰ Ä‘áº¡t 85.44% - sá»¥t giáº£m tá»›i gáº§n 9%! Táº¡i sao cÆ¡ cháº¿ Ä‘iá»u biáº¿n cá»§a nhÃ³m láº¡i máº¥t á»•n Ä‘á»‹nh vÃ  nháº¡y cáº£m cá»±c Ä‘oan theo kiáº¿n trÃºc máº¡ng nhÆ° váº­y? LÃ m sao Ä‘áº£m báº£o tÃ­nh an toÃ n khi triá»ƒn khai thá»±c táº¿?\"",
        [
            "Hai con sá»‘ trÃªn cho tháº¥y hiá»‡u quáº£ cá»§a FiLM phá»¥ thuá»™c vÃ o backbone; chÃºng khÃ´ng Ä‘á»§ Ä‘á»ƒ káº¿t luáº­n nguyÃªn nhÃ¢n lÃ  má»™t cÆ¡ cháº¿ kiáº¿n trÃºc cá»¥ thá»ƒ nhÆ° depthwise convolution hay compound scaling. FiLM trong dá»± Ã¡n sinh gamma vÃ  beta tá»« metadata rá»“i Ä‘iá»u biáº¿n feature áº£nh, nÃªn má»©c há»¯u Ã­ch cá»§a metadata phá»¥ thuá»™c vÃ o biá»ƒu diá»…n mÃ  tá»«ng backbone táº¡o ra vÃ  cÃ¡ch tá»‘i Æ°u há»™i tá»¥.",
            "VÃ¬ váº­y, nhÃ³m khÃ´ng kháº³ng Ä‘á»‹nh FiLM luÃ´n tá»‘t hÆ¡n hoáº·c á»•n Ä‘á»‹nh trÃªn má»i kiáº¿n trÃºc. Káº¿t luáº­n cÃ³ thá»ƒ báº£o vá»‡ lÃ : FiLM lÃ  má»™t cÆ¡ cháº¿ fusion cáº§n Ä‘Æ°á»£c tháº©m Ä‘á»‹nh theo tá»«ng backbone vÃ  tá»«ng táº­p dá»¯ liá»‡u; cáº¥u hÃ¬nh triá»ƒn khai chá»‰ Ä‘Æ°á»£c chá»n sau khi Ä‘Ã¡nh giÃ¡ trÃªn protocol cá»‘ Ä‘á»‹nh vÃ  bÃ¡o cÃ¡o Ä‘áº§y Ä‘á»§ AUC, sensitivity, specificity vÃ  Ä‘á»™ biáº¿n thiÃªn.",
            "Trong artefact EfficientNet-B4 hiá»‡n cÃ³ cá»§a dá»± Ã¡n, AUC HAM cá»§a image-only lÃ  85.37%, FiLM lÃ  85.43%, FiLM weighted lÃ  85.44%, cÃ²n gating lÃ  82.54%. Kiá»ƒm Ä‘á»‹nh AUC Ä‘ang lÆ°u khÃ´ng cho tháº¥y FiLM khÃ¡c image-only cÃ³ Ã½ nghÄ©a thá»‘ng kÃª. Do Ä‘Ã³, nhÃ³m khÃ´ng sá»­ dá»¥ng EfficientNet-B4 Ä‘á»ƒ kháº³ng Ä‘á»‹nh Æ°u tháº¿ phá»• quÃ¡t cá»§a FiLM.",
        ],
    ),
    (
        "3. BÃ¡o cÃ¡o Ä‘áº§y Ä‘á»§ bá»‘n chiáº¿n lÆ°á»£c fusion",
        "\"Trong bÃ¡o cÃ¡o vÃ  slide 4, nhÃ³m dÃ nh ráº¥t nhiá»u thá»i gian phÃ¢n tÃ­ch vÃ  Ä‘á» xuáº¥t Ä‘áº¿n 4 chiáº¿n lÆ°á»£c dung há»£p (Baseline, Concatenation, FiLM, Gating). Tháº¿ nhÆ°ng táº¡i sao trong báº£ng káº¿t quáº£ thá»±c nghiá»‡m, nhÃ³m chá»‰ Ä‘Æ°a ra duy nháº¥t káº¿t quáº£ cá»§a Strategy 3 (FiLM)? Káº¿t quáº£ cá»§a hai chiáº¿n lÆ°á»£c Concatenation vÃ  Gating Ä‘Ã¢u? Pháº£i chÄƒng chÃºng quÃ¡ tá»‡ nÃªn nhÃ³m Ä‘Ã£ giáº¥u Ä‘i?\"",
        [
            "Náº¿u má»™t slide chá»‰ trÃ¬nh bÃ y FiLM mÃ  khÃ´ng kÃ¨m báº£ng Ä‘á»‘i chá»©ng, Ä‘Ã³ lÃ  thiáº¿u sÃ³t trÃ¬nh bÃ y vÃ  cáº§n sá»­a. Trong code, nhÃ³m cÃ³ triá»ƒn khai image-only, concatenation, FiLM vÃ  gating; vÃ¬ váº­y bÃ¡o cÃ¡o khoa há»c pháº£i thá»ƒ hiá»‡n cÃ¡c káº¿t quáº£ Ä‘á»‘i chá»©ng thay vÃ¬ chá»‰ nÃªu cáº¥u hÃ¬nh tá»‘t nháº¥t.",
            "Vá»›i cÃ¡c checkpoint EfficientNet-B4 hiá»‡n Ä‘ang truy váº¿t Ä‘Æ°á»£c, HAM10000 cho AUC láº§n lÆ°á»£t: image-only 85.37%, FiLM 85.43%, FiLM weighted 85.44% vÃ  gating 82.54%. TrÃªn BCN20000, cÃ¡c giÃ¡ trá»‹ tÆ°Æ¡ng á»©ng lÃ  79.31%, 79.31%, 79.25% vÃ  79.22%. CÃ¡c sá»‘ liá»‡u nÃ y khÃ´ng há»— trá»£ tuyÃªn bá»‘ ráº±ng gating tÆ°Æ¡ng Ä‘Æ°Æ¡ng FiLM hoáº·c FiLM luÃ´n vÆ°á»£t trá»™i.",
            "Äá»‘i vá»›i Concatenation hoáº·c cÃ¡c backbone Ä‘Æ°á»£c nÃªu trong báº£ng cá»§a bÃ i bÃ¡o nhÆ°ng khÃ´ng cÃ³ Ä‘áº§y Ä‘á»§ artefact tÃ¡i láº­p trong workspace hiá»‡n táº¡i, nhÃ³m sáº½ Ä‘Ã­nh kÃ¨m báº£ng gá»‘c cÃ³ nguá»“n rÃµ rÃ ng hoáº·c cháº¡y láº¡i cÃ¹ng protocol trÆ°á»›c khi Ä‘Æ°a vÃ o káº¿t luáº­n. CÃ¡ch tráº£ lá»i Ä‘Ãºng lÃ  cÃ´ng khai toÃ n bá»™ so sÃ¡nh, khÃ´ng giáº£i thÃ­ch viá»‡c thiáº¿u báº£ng báº±ng suy Ä‘oÃ¡n vá» chi phÃ­ hay thá»i gian há»™i tá»¥.",
        ],
    ),
    (
        "4. Sensitivity trong bÃ i toÃ¡n sÃ ng lá»c",
        "\"Táº¡i sao trong báº£ng káº¿t quáº£ thá»±c nghiá»‡m, nhÃ³m chá»‰ bÃ¡o cÃ¡o hai chá»‰ sá»‘ lÃ  AUC vÃ  Äá»™ Ä‘áº·c hiá»‡u (Specificity)? Trong ung thÆ° da, bá» sÃ³t má»™t ca Ã¡c tÃ­nh, tá»©c Ã‚m tÃ­nh giáº£, má»›i lÃ  nguy cÆ¡ nghiÃªm trá»ng. Viá»‡c nhÃ³m khÃ´ng bÃ¡o cÃ¡o Äá»™ nháº¡y cÃ³ pháº£i lÃ  Ä‘á»ƒ che giáº¥u viá»‡c mÃ´ hÃ¬nh Ä‘ang bá» sÃ³t hÃ ng loáº¡t ca ung thÆ°?\"",
        [
            "Há»™i Ä‘á»“ng Ä‘áº·t Ä‘Ãºng trá»ng tÃ¢m: trong á»©ng dá»¥ng sÃ ng lá»c, sensitivity/recall pháº£i Ä‘Æ°á»£c bÃ¡o cÃ¡o cÃ¹ng AUC vÃ  specificity. Code Ä‘Ã¡nh giÃ¡ cá»§a nhÃ³m Ä‘Ã£ tÃ­nh recall, precision vÃ  F1, do Ä‘Ã³ viá»‡c khÃ´ng Ä‘Æ°a sensitivity lÃªn slide lÃ  thiáº¿u sÃ³t bÃ¡o cÃ¡o chá»© khÃ´ng pháº£i má»™t chá»‰ sá»‘ khÃ´ng cÃ³ sáºµn.",
            "Quan trá»ng hÆ¡n, káº¿t quáº£ EfficientNet-B4 hiá»‡n cÃ³ cho HAM10000 táº¡i threshold 0.5 cho recall khÃ¡ tháº¥p: image-only 36.83%, FiLM 40.40%, FiLM weighted 37.52% vÃ  gating 23.66%. TrÃªn BCN20000, recall láº§n lÆ°á»£t lÃ  68.75%, 70.35%, 70.18% vÃ  68.41%. NhÃ³m khÃ´ng thá»ƒ dÃ¹ng AUC hoáº·c specificity Ä‘á»ƒ kháº³ng Ä‘á»‹nh há»‡ thá»‘ng Ä‘Ã£ an toÃ n cho sÃ ng lá»c lÃ¢m sÃ ng.",
            "Káº¿t luáº­n phÃ¹ há»£p lÃ  mÃ´ hÃ¬nh hiá»‡n má»›i á»Ÿ má»©c nghiÃªn cá»©u phÃ¢n loáº¡i nguy cÆ¡. TrÆ°á»›c khi hÆ°á»›ng tá»›i sá»­ dá»¥ng lÃ¢m sÃ ng, nhÃ³m pháº£i lá»±a chá»n threshold theo má»¥c tiÃªu sensitivity, bÃ¡o cÃ¡o confusion matrix vÃ  khoáº£ng tin cáº­y, Ä‘á»“ng thá»i Ä‘Ã¡nh giÃ¡ Ä‘Ã¡nh Ä‘á»•i giá»¯a giáº£m false negative vÃ  tÄƒng false positive trÃªn táº­p kiá»ƒm Ä‘á»‹nh Ä‘á»™c láº­p.",
        ],
    ),
    (
        "5. DullRazor, CLAHE vÃ  thÃ´ng tin mÃ u sáº¯c",
        "\"NhÃ³m sá»­ dá»¥ng thuáº­t toÃ¡n DullRazor Ä‘á»ƒ xÃ³a lÃ´ng vÃ  thuáº­t toÃ¡n CLAHE Ä‘á»ƒ tÄƒng cÆ°á»ng tÆ°Æ¡ng pháº£n. Tuy nhiÃªn, trong y khoa, quy trÃ¬nh ABCD Ä‘á»ƒ cháº©n Ä‘oÃ¡n u háº¯c tá»‘ dá»±a ráº¥t máº¡nh vÃ o yáº¿u tá»‘ C (Color). Thuáº­t toÃ¡n CLAHE lÃ m biáº¿n Ä‘á»•i cá»¥c bá»™ biá»ƒu Ä‘á»“ mÃ u sáº¯c vÃ  Ä‘á»™ sÃ¡ng cá»§a bá»©c áº£nh. Cháº³ng pháº£i nhÃ³m Ä‘ang tá»± tay phÃ¡ há»§y cÃ¡c Ä‘áº·c trÆ°ng mÃ u sáº¯c sinh há»c cá»‘t lÃµi, Ã©p AI há»c trÃªn má»™t bá»©c áº£nh bá»‹ bÃ³p mÃ©o hay sao?\"",
        [
            "ÄÃ¢y lÃ  rá»§i ro há»£p lá»‡ vÃ  nhÃ³m khÃ´ng kháº³ng Ä‘á»‹nh áº£nh sau tiá»n xá»­ lÃ½ giá»¯ nguyÃªn mÃ u theo nghÄ©a pixel-by-pixel. Trong code, bÆ°á»›c CLAHE Ä‘Æ°á»£c thá»±c hiá»‡n sau khi chuyá»ƒn áº£nh sang LAB vÃ  chá»‰ Ã¡p dá»¥ng trÃªn kÃªnh L vá»›i clipLimit báº±ng 1.5; hai kÃªnh sáº¯c Ä‘á»™ a vÃ  b khÃ´ng bá»‹ CLAHE biáº¿n Ä‘á»•i trá»±c tiáº¿p. Má»¥c tiÃªu cá»§a bÆ°á»›c nÃ y lÃ  giáº£m biáº¿n thiÃªn Ä‘á»™ sÃ¡ng vÃ  lÃ m rÃµ cáº¥u trÃºc quan sÃ¡t Ä‘Æ°á»£c.",
            "Tuy nhiÃªn, pipeline Ä‘áº§y Ä‘á»§ cÃ²n cÃ³ Gray-World trÆ°á»›c CLAHE Ä‘á»ƒ hiá»‡u chá»‰nh cÃ¢n báº±ng mÃ u, inpainting táº¡i vÃ¹ng Ä‘Æ°á»£c nháº­n diá»‡n lÃ  lÃ´ng, vÃ  bilateral filtering sau khi tÄƒng tÆ°Æ¡ng pháº£n. NgoÃ i ra, táº­p train cÃ³ ColorJitter. VÃ¬ váº­y, cÃ¢u nÃ³i ráº±ng cÃ¡c Ä‘áº·c trÆ°ng mÃ u sinh há»c Ä‘Æ°á»£c báº£o toÃ n nguyÃªn váº¹n lÃ  khÃ´ng chÃ­nh xÃ¡c.",
            "Láº­p luáº­n khoa há»c Ä‘Ãºng lÃ  pipeline cá»‘ gáº¯ng giáº£m nhiá»…u do lÃ´ng vÃ  chiáº¿u sÃ¡ng trong khi háº¡n cháº¿ can thiá»‡p trá»±c tiáº¿p vÃ o sáº¯c Ä‘á»™ á»Ÿ bÆ°á»›c CLAHE. Äá»ƒ xÃ¡c nháº­n lá»±a chá»n nÃ y khÃ´ng lÃ m máº¥t tÃ­n hiá»‡u cháº©n Ä‘oÃ¡n, nhÃ³m cáº§n ablation áº£nh gá»‘c, xÃ³a lÃ´ng, Gray-World vÃ  CLAHE; Ä‘á»“ng thá»i so sÃ¡nh sensitivity/AUC, Ä‘áº·c biá»‡t vá»›i melanoma, trÆ°á»›c khi káº¿t luáº­n lá»£i Ã­ch cá»§a tiá»n xá»­ lÃ½.",
        ],
    ),
    (
        "6. Giá»›i háº¡n cá»§a Grad-CAM vÃ  SHAP",
        "\"NhÃ³m sá»­ dá»¥ng Grad-CAM Ä‘á»ƒ chá»©ng minh AI Ä‘Ã£ 'nhÃ¬n' vÃ o Ä‘Ãºng chá»— tá»•n thÆ°Æ¡ng nháº±m xÃ¢y dá»±ng niá»m tin lÃ¢m sÃ ng. Tuy nhiÃªn, Grad-CAM cÃ³ thá»ƒ bá»‹ áº£nh hÆ°á»Ÿng bá»Ÿi vÃ¹ng gradient máº¡nh hoáº·c tÆ°Æ¡ng pháº£n do CLAHE táº¡o ra. LÃ m sao nhÃ³m kháº³ng Ä‘á»‹nh vÃ¹ng mÃ u Ä‘á» thá»±c sá»± lÃ  tri thá»©c y khoa?\"",
        [
            "NhÃ³m khÃ´ng nÃªn dÃ¹ng tá»« 'chá»©ng minh' cho Grad-CAM. Grad-CAM lÃ  cÃ´ng cá»¥ trá»±c quan hÃ³a vÃ¹ng áº£nh áº£nh hÆ°á»Ÿng tá»›i dá»± Ä‘oÃ¡n, khÃ´ng pháº£i xÃ¡c nháº­n ráº±ng mÃ´ hÃ¬nh Ä‘Ã£ há»c Ä‘Ãºng Ä‘áº·c trÆ°ng y khoa hoáº·c ráº±ng vÃ¹ng chÃº Ã½ khÃ´ng chá»‹u áº£nh hÆ°á»Ÿng cá»§a preprocessing.",
            "Trong triá»ƒn khai hiá»‡n táº¡i, hÃ m sinh Grad-CAM Ä‘Æ°a metadata cá»‘ Ä‘á»‹nh báº±ng tensor zero khi gá»i mÃ´ hÃ¬nh. VÃ¬ váº­y, cÃ¡c heatmap hiá»‡n táº¡i chá»§ yáº¿u kháº£o sÃ¡t nhÃ¡nh áº£nh dÆ°á»›i má»™t Ä‘iá»u kiá»‡n metadata cá»‘ Ä‘á»‹nh; chÃºng chÆ°a chá»©ng minh sá»± tÆ°Æ¡ng tÃ¡c Ä‘á»™ng giá»¯a áº£nh vÃ  metadata. CÃ¡c script SHAP cÅ©ng lÃ  phÃ¢n tÃ­ch Ä‘Ã³ng gÃ³p metadata, khÃ´ng pháº£i báº±ng chá»©ng nhÃ¢n quáº£ vá» vá»‹ trÃ­ tá»•n thÆ°Æ¡ng.",
            "CÃ¡ch tráº£ lá»i phÃ¹ há»£p lÃ : Grad-CAM vÃ  SHAP Ä‘Æ°á»£c sá»­ dá»¥ng nhÆ° phÃ¢n tÃ­ch há»— trá»£. Äá»ƒ kiá»ƒm tra nghiÃªm ngáº·t lo ngáº¡i cá»§a há»™i Ä‘á»“ng, nhÃ³m cáº§n bá»• sung so sÃ¡nh heatmap trÃªn áº£nh gá»‘c vÃ  áº£nh tiá»n xá»­ lÃ½, kiá»ƒm thá»­ occlusion/perturbation Ä‘á»‹nh lÆ°á»£ng, vÃ  náº¿u Ä‘Ã¡nh giÃ¡ multimodal saliency thÃ¬ pháº£i giá»¯ metadata thá»±c hoáº·c thay Ä‘á»•i metadata cÃ³ kiá»ƒm soÃ¡t.",
        ],
    ),
    (
        "7. Kháº£ nÄƒng triá»ƒn khai trÃªn thiáº¿t bá»‹ di Ä‘á»™ng",
        "\"NhÃ³m Ä‘á»‹nh hÆ°á»›ng nÃ©n mÃ´ hÃ¬nh Ä‘Æ°a lÃªn á»©ng dá»¥ng di Ä‘á»™ng, Ä‘á»“ng thá»i tÃ­ch há»£p LLM Ä‘á»ƒ Ä‘á»c EHR. Má»™t Vision Transformer káº¿t há»£p LLM Ä‘Ã²i há»i tÃ i nguyÃªn lá»›n. LÃ m sao cháº¡y trÃªn Ä‘iá»‡n thoáº¡i cáº¥u hÃ¬nh tháº¥p á»Ÿ tráº¡m y táº¿ xÃ£?\"",
        [
            "ÄÃ¢y lÃ  Ä‘á»‹nh hÆ°á»›ng tÆ°Æ¡ng lai, chÆ°a pháº£i nÄƒng lá»±c Ä‘Ã£ Ä‘Æ°á»£c triá»ƒn khai hoáº·c benchmark trong dá»± Ã¡n hiá»‡n táº¡i. NhÃ³m khÃ´ng chá»§ trÆ°Æ¡ng cháº¡y Ä‘á»“ng thá»i má»™t ViT lá»›n vÃ  má»™t LLM y táº¿ hoÃ n chá»‰nh offline trÃªn Ä‘iá»‡n thoáº¡i cáº¥u hÃ¬nh tháº¥p.",
            "Kiáº¿n trÃºc kháº£ thi vá» máº·t thiáº¿t káº¿ lÃ  phÃ¢n tÃ¡ch nhiá»‡m vá»¥: thiáº¿t bá»‹ Ä‘áº§u cuá»‘i cÃ³ thá»ƒ cháº¡y má»™t bá»™ mÃ£ hÃ³a áº£nh nhá» sau khi Ä‘Ã£ Ä‘Æ°á»£c lá»±a chá»n, lÆ°á»£ng tá»­ hÃ³a hoáº·c distillation; tÃ¡c vá»¥ xá»­ lÃ½ vÄƒn báº£n EHR phá»©c táº¡p chá»‰ thá»±c hiá»‡n trÃªn mÃ¡y chá»§ khi cÃ³ háº¡ táº§ng máº¡ng, kiá»ƒm soÃ¡t truy cáº­p vÃ  báº£o máº­t dá»¯ liá»‡u phÃ¹ há»£p.",
            "Tuy nhiÃªn, Ä‘á»ƒ biáº¿n Ä‘á»‹nh hÆ°á»›ng thÃ nh káº¿t quáº£ nghiÃªn cá»©u, nhÃ³m pháº£i bÃ¡o cÃ¡o kÃ­ch thÆ°á»›c model, latency, bá»™ nhá»›, má»©c giáº£m hiá»‡u nÄƒng sau nÃ©n, kháº£ nÄƒng hoáº¡t Ä‘á»™ng offline vÃ  quy trÃ¬nh báº£o máº­t. Hiá»‡n táº¡i nhÃ³m chá»‰ trÃ¬nh bÃ y Ä‘Ã¢y lÃ  lá»™ trÃ¬nh triá»ƒn khai, khÃ´ng pháº£i káº¿t quáº£ thá»±c nghiá»‡m Ä‘Ã£ chá»©ng minh.",
        ],
    ),
    (
        "8. Ã nghÄ©a khoa há»c cá»§a má»©c tÄƒng AUC",
        "\"Strategy 3 (FiLM) Ä‘áº¡t AUC 94.02%, cao hÆ¡n Baseline khoáº£ng 1.5 - 2%. LÃ m sao nhÃ³m chá»©ng minh Ä‘Ã¢y lÃ  sá»± vÆ°á»£t trá»™i cÃ³ Ã½ nghÄ©a khoa há»c chá»© khÃ´ng pháº£i do nhiá»…u ngáº«u nhiÃªn hoáº·c khá»Ÿi táº¡o trá»ng sá»‘?\"",
        [
            "Náº¿u con sá»‘ 94.02% Ä‘Æ°á»£c láº¥y tá»« báº£ng ConvNeXt cá»§a bÃ i bÃ¡o, nhÃ³m chá»‰ cÃ³ thá»ƒ nÃ³i Ä‘Ã³ lÃ  káº¿t quáº£ trung bÃ¬nh theo cÃ¡c fold Ä‘Æ°á»£c bÃ¡o cÃ¡o trong báº£ng tÆ°Æ¡ng á»©ng. KhÃ´ng Ä‘Æ°á»£c tá»± Ä‘á»™ng diá»…n giáº£i thÃ nh káº¿t quáº£ cá»§a nhiá»u random seed náº¿u chÆ°a cÃ³ thÃ­ nghiá»‡m seed robustness riÃªng.",
            "Trong workspace hiá»‡n táº¡i, kiá»ƒm Ä‘á»‹nh paired theo 5 fold Ä‘Ã£ Ä‘Æ°á»£c lÆ°u cho EfficientNet-B4. TrÃªn HAM10000, FiLM weighted chá»‰ cao hÆ¡n image-only khoáº£ng 0.075 Ä‘iá»ƒm pháº§n trÄƒm vá» AUC vá»›i p = 0.8531; trÃªn BCN20000, FiLM weighted tháº¥p hÆ¡n image-only khoáº£ng 0.062 Ä‘iá»ƒm pháº§n trÄƒm vá»›i p = 0.8920. Nhá»¯ng káº¿t quáº£ nÃ y khÃ´ng cho phÃ©p tuyÃªn bá»‘ FiLM vÆ°á»£t trá»™i cÃ³ Ã½ nghÄ©a thá»‘ng kÃª trÃªn EfficientNet-B4.",
            "Äá»ƒ báº£o vá»‡ tuyÃªn bá»‘ cho ConvNeXt 94.02%, nhÃ³m cáº§n cung cáº¥p dá»± Ä‘oÃ¡n theo fold hoáº·c theo máº«u cá»§a chÃ­nh cáº¥u hÃ¬nh ConvNeXt, thá»±c hiá»‡n kiá»ƒm Ä‘á»‹nh phÃ¹ há»£p vÃ  bÃ¡o cÃ¡o khoáº£ng tin cáº­y. Cho Ä‘áº¿n khi cÃ³ kiá»ƒm Ä‘á»‹nh tÆ°Æ¡ng á»©ng, káº¿t luáº­n nÃªn giá»›i háº¡n á»Ÿ má»©c: FiLM cho káº¿t quáº£ triá»ƒn vá»ng á»Ÿ má»™t sá»‘ backbone trong báº£ng bÃ¡o cÃ¡o, nhÆ°ng má»©c cáº£i thiá»‡n chÆ°a Ä‘Æ°á»£c chá»©ng minh lÃ  phá»• quÃ¡t.",
        ],
    ),
    (
        "9. Focal Loss vÃ  hiá»‡u chuáº©n xÃ¡c suáº¥t",
        "\"NhÃ³m dÃ¹ng Focal Loss Ä‘á»ƒ xá»­ lÃ½ máº¥t cÃ¢n báº±ng. HÃ m nÃ y cÃ³ thá»ƒ lÃ m sai lá»‡ch confidence score. DÃ¹ Accuracy cao, Ä‘á»™ tin cáº­y cá»§a xÃ¡c suáº¥t Ä‘Æ°a cho bÃ¡c sÄ© cÃ³ thá»ƒ bá»‹ sai. NhÃ³m Ä‘Ã£ Ä‘o lÆ°á»ng ECE chÆ°a?\"",
        [
            "NhÃ³m thá»«a nháº­n Ä‘iá»ƒm nÃ y. Dá»± Ã¡n hiá»‡n sá»­ dá»¥ng Focal BCE Loss vá»›i alpha báº±ng 0.75 vÃ  gamma báº±ng 2.0 Ä‘á»ƒ táº­p trung há»c cÃ¡c máº«u khÃ³, nhÆ°ng trong code vÃ  artefact hiá»‡n cÃ³ chÆ°a cÃ³ Ä‘Ã¡nh giÃ¡ ECE, Brier score hoáº·c calibration curve.",
            "Do Ä‘Ã³, Ä‘áº§u ra sigmoid hiá»‡n táº¡i khÃ´ng nÃªn Ä‘Æ°á»£c mÃ´ táº£ nhÆ° má»™t xÃ¡c suáº¥t nguy cÆ¡ Ä‘Ã£ Ä‘Æ°á»£c hiá»‡u chuáº©n Ä‘á»ƒ bÃ¡c sÄ© sá»­ dá»¥ng trá»±c tiáº¿p. AUC Ä‘Ã¡nh giÃ¡ kháº£ nÄƒng xáº¿p háº¡ng vÃ  cÃ¡c chá»‰ sá»‘ táº¡i threshold Ä‘Ã¡nh giÃ¡ quyáº¿t Ä‘á»‹nh phÃ¢n lá»›p; chÃºng khÃ´ng chá»©ng minh confidence score Ä‘Ã¡ng tin cáº­y.",
            "BÆ°á»›c bá»• sung báº¯t buá»™c trÆ°á»›c khi trÃ¬nh bÃ y xÃ¡c suáº¥t lÃ¢m sÃ ng lÃ  hiá»‡u chuáº©n trÃªn táº­p validation tÃ¡ch biá»‡t, cháº³ng háº¡n temperature scaling, sau Ä‘Ã³ bÃ¡o cÃ¡o ECE, Brier score, reliability diagram vÃ  Ä‘Ã¡nh giÃ¡ láº¡i sensitivity/specificity táº¡i threshold Ä‘Æ°á»£c chá»n.",
        ],
    ),
    (
        "10. Lá»£i Ã­ch metadata hay chá»‰ do tÄƒng sá»‘ tham sá»‘",
        "\"FiLM cho káº¿t quáº£ cao hÆ¡n Baseline, nhÆ°ng FiLM cÅ©ng thÃªm tham sá»‘. LÃ m sao nhÃ³m biáº¿t má»©c tÄƒng Ä‘áº¿n tá»« thÃ´ng tin metadata thay vÃ¬ capacity lá»›n hÆ¡n giÃºp mÃ´ hÃ¬nh overfit?\"",
        [
            "So sÃ¡nh image-only vá»›i FiLM cÃ³ metadata tháº­t chÆ°a Ä‘á»§ Ä‘á»ƒ tÃ¡ch riÃªng hai giáº£ thuyáº¿t: lá»£i Ã­ch cá»§a thÃ´ng tin lÃ¢m sÃ ng vÃ  lá»£i Ã­ch cá»§a viá»‡c tÄƒng sá»‘ tham sá»‘. Trong workspace hiá»‡n táº¡i, nhÃ³m chÆ°a cÃ³ artefact chá»©ng minh Ä‘Ã£ cháº¡y thÃ­ nghiá»‡m shuffled metadata hoáº·c dummy metadata.",
            "ThÃ­ nghiá»‡m Ä‘Ãºng Ä‘á»ƒ tráº£ lá»i cÃ¢u há»i lÃ  giá»¯ nguyÃªn kiáº¿n trÃºc FiLM vÃ  sá»‘ tham sá»‘, nhÆ°ng huáº¥n luyá»‡n/Ä‘Ã¡nh giÃ¡ vá»›i metadata bá»‹ hoÃ¡n vá»‹ giá»¯a cÃ¡c máº«u hoáº·c metadata háº±ng sá»‘. Náº¿u FiLM vá»›i metadata tháº­t vÆ°á»£t rÃµ rá»‡t cÃ¡c Ä‘á»‘i chá»©ng cÃ¹ng capacity, khi Ä‘Ã³ má»›i cÃ³ báº±ng chá»©ng ráº±ng thÃ´ng tin lÃ¢m sÃ ng Ä‘Ã³ng gÃ³p thá»±c sá»±.",
            "VÃ¬ chÆ°a cÃ³ ablation Ä‘Ã³ trong káº¿t quáº£ hiá»‡n táº¡i, nhÃ³m chá»‰ káº¿t luáº­n ráº±ng mÃ´ hÃ¬nh Ä‘Ã£ triá»ƒn khai cÆ¡ cháº¿ Ä‘iá»u biáº¿n báº±ng metadata vÃ  cÃ³ thá»ƒ kháº£o sÃ¡t Ä‘Ã³ng gÃ³p feature; nhÃ³m chÆ°a tuyÃªn bá»‘ má»©c tÄƒng AUC, náº¿u cÃ³, hoÃ n toÃ n do Ã½ nghÄ©a y khoa cá»§a metadata.",
        ],
    ),
    (
        "11. Domain shift vÃ  phÃ©p so sÃ¡nh giá»¯a HAM10000/BCN20000",
        "\"MÃ´ hÃ¬nh giáº£m tá»« khoáº£ng 94% trÃªn HAM10000 xuá»‘ng 84.88% trÃªn BCN20000. Viá»‡c giáº£m Ä‘áº¿n 10% chá»©ng tá» mÃ´ hÃ¬nh bá»‹ Domain Shift ráº¥t náº·ng. MÃ´ hÃ¬nh cÃ³ thá»±c sá»± há»c Ä‘Æ°á»£c khÃ¡i quÃ¡t sinh há»c khÃ´ng?\"",
        [
            "Cáº§n chá»‰nh láº¡i tiá»n Ä‘á» cá»§a phÃ©p so sÃ¡nh. Hai sá»‘ 94.02% trÃªn HAM10000 vÃ  84.88% trÃªn BCN20000 trong tÃ i liá»‡u trÃ¬nh bÃ y gáº¯n vá»›i cÃ¡c cáº¥u hÃ¬nh/backbone bÃ¡o cÃ¡o khÃ¡c nhau; so sÃ¡nh trá»±c tiáº¿p chÃºng khÃ´ng pháº£i lÃ  má»™t phÃ©p Ä‘o domain shift há»£p lá»‡.",
            "ÄÃ¡nh giÃ¡ domain shift Ä‘Ãºng pháº£i cá»‘ Ä‘á»‹nh mÃ´ hÃ¬nh vÃ  hÆ°á»›ng chuyá»ƒn miá»n: vÃ­ dá»¥ train trÃªn HAM rá»“i test trá»±c tiáº¿p trÃªn BCN, hoáº·c ngÆ°á»£c láº¡i, vá»›i cÃ¹ng checkpoint, cÃ¹ng Ã¡nh xáº¡ metadata vÃ  cÃ¹ng chá»‰ sá»‘. Dá»± Ã¡n cÃ³ script cho Ä‘Ã¡nh giÃ¡ chÃ©o, nhÆ°ng trong artefact hiá»‡n Ä‘ang Ä‘á»‘i chiáº¿u chÆ°a cÃ³ báº£ng káº¿t quáº£ cross-dataset Ä‘Æ°á»£c lÆ°u Ä‘á»ƒ dÃ¹ng lÃ m báº±ng chá»©ng Ä‘á»‹nh lÆ°á»£ng.",
            "VÃ¬ váº­y, nhÃ³m chá»‰ cÃ³ thá»ƒ nÃ³i hai bá»™ dá»¯ liá»‡u cÃ³ khÃ¡c biá»‡t phÃ¢n phá»‘i vÃ  viá»‡c Ä‘Ã¡nh giÃ¡ chÃ©o lÃ  cáº§n thiáº¿t. NhÃ³m chÆ°a dÃ¹ng chÃªnh lá»‡ch giá»¯a hai báº£ng ná»™i miá»n Ä‘á»ƒ tuyÃªn bá»‘ mÃ´ hÃ¬nh Ä‘Ã£ khÃ¡ng domain shift hoáº·c Ä‘Ã£ há»c Ä‘Æ°á»£c khÃ¡i quÃ¡t sinh há»c.",
        ],
    ),
    (
        "12. Biáº¿n metadata nÃ o Ä‘Ã³ng gÃ³p chÃ­nh",
        "\"NhÃ³m náº¡p Tuá»•i, Giá»›i tÃ­nh, Vá»‹ trÃ­ vÃ o mÃ´ hÃ¬nh. Trong má»©c tÄƒng hiá»‡u suáº¥t Ä‘Ã³, biáº¿n nÃ o Ä‘Ã³ng vai trÃ² chÃ­nh? Náº¿u loáº¡i bá» Giá»›i tÃ­nh, mÃ´ hÃ¬nh cÃ³ bá»‹ áº£nh hÆ°á»Ÿng khÃ´ng?\"",
        [
            "Trong dá»¯ liá»‡u HAM, nhÃ¡nh metadata sá»­ dá»¥ng age, sex vÃ  localization; trong BCN sá»­ dá»¥ng age_approx, sex vÃ  anatom_site_general. CÃ¡c báº£ng importance hiá»‡n cÃ³ cho tháº¥y tuá»•i lÃ  feature cÃ³ trá»ng sá»‘ giáº£i thÃ­ch lá»›n nháº¥t: khoáº£ng 0.630 trong artefact HAM vÃ  khoáº£ng 0.807 trong artefact BCN; sau Ä‘Ã³ lÃ  cÃ¡c nhÃ³m vá»‹ trÃ­ giáº£i pháº«u, cÃ²n cÃ¡c biáº¿n giá»›i tÃ­nh cÃ³ importance nhá» hÆ¡n.",
            "Tuy nhiÃªn, feature importance hoáº·c SHAP chá»‰ pháº£n Ã¡nh má»©c Ä‘Ã³ng gÃ³p trong mÃ´ hÃ¬nh/phÃ¢n tÃ­ch hiá»‡n táº¡i; chÃºng khÃ´ng tÆ°Æ¡ng Ä‘Æ°Æ¡ng vá»›i báº±ng chá»©ng ráº±ng bá» giá»›i tÃ­nh chá»‰ lÃ m AUC thay Ä‘á»•i má»™t con sá»‘ xÃ¡c Ä‘á»‹nh. Muá»‘n tráº£ lá»i Ä‘á»‹nh lÆ°á»£ng, nhÃ³m pháº£i cháº¡y ablation loáº¡i riÃªng age, sex vÃ  location trÃªn cÃ¹ng protocol.",
            "Do Ä‘Ã³, káº¿t luáº­n Ä‘Æ°á»£c phÃ©p nÃªu lÃ : phÃ¢n tÃ­ch hiá»‡n cÃ³ gá»£i Ã½ tuá»•i lÃ  tÃ­n hiá»‡u metadata ná»•i báº­t nháº¥t vÃ  giá»›i tÃ­nh cÃ³ Ä‘Ã³ng gÃ³p nhá» hÆ¡n trong cÃ¡c artefact Ä‘ang lÆ°u. NhÃ³m chÆ°a kháº³ng Ä‘á»‹nh tÃ¡c Ä‘á»™ng nhÃ¢n quáº£ cá»§a tá»«ng biáº¿n lÃªn AUC khi chÆ°a cÃ³ thÃ­ nghiá»‡m loáº¡i biáº¿n.",
        ],
    ),
    (
        "13. Thiáº¿u hoáº·c nháº­p sai metadata khi triá»ƒn khai",
        "\"FiLM giÃºp mÃ´ hÃ¬nh chÃº Ã½ metadata, nhÆ°ng náº¿u bÃ¡c sÄ© nháº­p sai tuá»•i hoáº·c quÃªn nháº­p vá»‹ trÃ­ thÃ¬ sao? MÃ´ hÃ¬nh cÃ³ phá»¥ thuá»™c quÃ¡ má»©c vÃ o metadata vÃ  sá»¥p Ä‘á»• so vá»›i image-only khÃ´ng?\"",
        [
            "ÄÃ¢y lÃ  rá»§i ro triá»ƒn khai cáº§n Ä‘Ã¡nh giÃ¡ riÃªng. Code hiá»‡n táº¡i há»— trá»£ giÃ¡ trá»‹ phÃ¢n loáº¡i unknown trong xá»­ lÃ½ dá»¯ liá»‡u vÃ  FiLM Ä‘Æ°á»£c khá»Ÿi táº¡o ban Ä‘áº§u gáº§n identity do lá»›p sinh gamma/beta khá»Ÿi táº¡o zero. Tuy nhiÃªn, sau huáº¥n luyá»‡n, Ä‘iá»u Ä‘Ã³ khÃ´ng báº£o Ä‘áº£m mÃ´ hÃ¬nh tá»± Ä‘á»™ng trá»Ÿ vá» image-only khi metadata thiáº¿u hoáº·c sai.",
            "Trong artefact hiá»‡n cÃ³, nhÃ³m chÆ°a cÃ³ robustness test che metadata, nhiá»…u tuá»•i, hoÃ¡n Ä‘á»•i vá»‹ trÃ­ hoáº·c Ä‘o má»©c suy giáº£m hiá»‡u nÄƒng khi nháº­p sai. VÃ¬ váº­y, nhÃ³m khÃ´ng tuyÃªn bá»‘ mÃ´ hÃ¬nh cÃ³ cÆ¡ cháº¿ safe-fail Ä‘Ã£ Ä‘Æ°á»£c chá»©ng minh.",
            "TrÆ°á»›c triá»ƒn khai, nhÃ³m cáº§n bá»• sung missing-modality/mis-entry testing, huáº¥n luyá»‡n vá»›i metadata dropout hoáº·c missingness mask, vÃ  cÃ¢n nháº¯c cung cáº¥p má»™t nhÃ¡nh image-only fallback Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ Ä‘á»™c láº­p. Khi Ä‘Ã³ má»›i cÃ³ thá»ƒ quy Ä‘á»‹nh rÃµ há»‡ thá»‘ng xá»­ lÃ½ trÆ°á»ng há»£p metadata khÃ´ng Ä‘Ã¡ng tin cáº­y nhÆ° tháº¿ nÃ o.",
        ],
    ),
]


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(document, text, *, size=11, bold=False, color=None, space_after=4):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Bá»˜ CÃ‚U Há»ŽI PHáº¢N BIá»†N VÃ€ TRáº¢ Lá»œI ÄÃƒ HIá»†U CHá»ˆNH")
    set_font(run, size=15, bold=True, color=(31, 78, 121))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Äá»‘i chiáº¿u vá»›i code vÃ  artefact hiá»‡n cÃ³ cá»§a dá»± Ã¡n cháº©n Ä‘oÃ¡n ung thÆ° da Ä‘a phÆ°Æ¡ng thá»©c")
    set_font(run, size=11, color=(89, 89, 89))

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(12)
    note.paragraph_format.line_spacing = 1.15
    run = note.add_run(
        "NguyÃªn táº¯c hiá»‡u chá»‰nh: chá»‰ phÃ¡t biá»ƒu cÃ¡c káº¿t luáº­n cÃ³ thá»ƒ kiá»ƒm chá»©ng tá»« code hoáº·c káº¿t quáº£ Ä‘ang lÆ°u; "
        "cÃ¡c thÃ­ nghiá»‡m chÆ°a cÃ³ artefact Ä‘Æ°á»£c nÃªu rÃµ lÃ  giá»›i háº¡n hoáº·c cÃ´ng viá»‡c cáº§n bá»• sung."
    )
    set_font(run, size=10, bold=True, color=(192, 0, 0))

    for title_text, question, answers in QA_ITEMS:
        add_paragraph(document, title_text, size=12, bold=True, color=(31, 78, 121), space_after=3)
        add_paragraph(document, question, size=11, bold=True, space_after=5)
        add_paragraph(document, "Tráº£ lá»i Ä‘á» xuáº¥t:", size=11, bold=True, color=(0, 97, 0), space_after=3)
        for answer in answers:
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.first_line_indent = Cm(-0.35)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.15
            bullet = paragraph.add_run("- ")
            set_font(bullet, size=11, bold=True)
            run = paragraph.add_run(answer)
            set_font(run, size=11)
        document.add_paragraph()

    add_paragraph(document, "CÄƒn cá»© Ä‘á»‘i chiáº¿u chÃ­nh trong workspace", size=12, bold=True, color=(31, 78, 121))
    references = [
        "src/preprocessed/preprocess_pipeline.py: DullRazor-style inpainting, Gray-World, CLAHE trÃªn kÃªnh L cá»§a LAB vÃ  bilateral filtering.",
        "src/data_logic/common_transforms.py: ColorJitter chá»‰ Ã¡p dá»¥ng trong augmentation cá»§a táº­p train.",
        "src/utils/experiment_runner.py vÃ  script/data/create_group_splits.py: protocol nhÃ³m theo lesion_id.",
        "src/models/fusion_head.py vÃ  src/models/__init__.py: triá»ƒn khai Concatenation, FiLM vÃ  Gating.",
        "src/utils/trainer.py: Ä‘Ã¡nh giÃ¡ AUC, accuracy, F1, precision, recall vÃ  specificity; Grad-CAM sá»­ dá»¥ng metadata cá»‘ Ä‘á»‹nh báº±ng zero tensor.",
        "results/significance_tests_auc.csv vÃ  cÃ¡c tá»‡p summary trong checkpoint_ham10000/checkpoint_bcn20000: sá»‘ liá»‡u EfficientNet-B4 vÃ  kiá»ƒm Ä‘á»‹nh AUC hiá»‡n cÃ³.",
    ]
    for reference in references:
        add_paragraph(document, "- " + reference, size=10, space_after=3)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH.resolve())
